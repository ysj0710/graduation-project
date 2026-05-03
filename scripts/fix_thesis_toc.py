from __future__ import annotations

from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/Users/enjoy0710/Desktop/学校相关/毕业设计/graduation-project")
DOCX = ROOT / "基于Node_js_Vue的个人财务记账与可视化分析系统设计与实现-完善稿.docx"
DOWNLOADS_DOCX = Path("/Users/enjoy0710/Downloads/基于Node_js_Vue的个人财务记账与可视化分析系统设计与实现-完善稿.docx")


def set_run_font(run, size=10.5):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def ensure_toc_styles(doc: Document):
    # WPS模板中已有 toc 2/toc 3，但通常没有 toc 4；这里补齐并统一缩进。
    for name, bold in [("toc 2", True), ("toc 3", False), ("toc 4", False)]:
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
        style.font.bold = bold
        style.paragraph_format.left_indent = Cm(0)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.tab_stops.clear_all()
        style.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def set_heading_outline_levels(doc: Document):
    # 当前论文模板用 Heading 2/3/4 作为章、节、小节；显式设置大纲级别，保证WPS按1-4级目录更新时包含1.2.1这类标题。
    for name, level in [("Heading 2", "1"), ("Heading 3", "2"), ("Heading 4", "3")]:
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        ppr = style.element.get_or_add_pPr()
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            ppr.append(outline)
        outline.set(qn("w:val"), level)


def add_update_fields_setting(doc: Document):
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)


def fld_char(kind: str):
    r = OxmlElement("w:r")
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), kind)
    if kind == "begin":
        fld.set(qn("w:dirty"), "true")
    r.append(fld)
    return r


def instr_text(text: str):
    r = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = text
    r.append(instr)
    return r


def collect_entries(doc: Document, toc_heading_idx: int):
    entries = []
    abstract_seen = False
    for idx, p in enumerate(doc.paragraphs):
        text = " ".join(p.text.split())
        if not text or text == "目 录":
            continue
        style = p.style.name
        if text == "摘 要":
            entries.append((2, text, "I", idx))
        elif text == "Abstract" and not abstract_seen:
            entries.append((2, text, "II", idx))
            abstract_seen = True
        elif style == "Heading 2" and idx > toc_heading_idx:
            entries.append((2, text, None, idx))
        elif style == "Heading 3" and idx > toc_heading_idx:
            entries.append((3, text, None, idx))
        elif style == "Heading 4" and idx > toc_heading_idx:
            entries.append((4, text, None, idx))
    # 目录中不放正文前的“目 录”本身；绪论以后页码由WPS更新域时重算。
    content_indices = [idx for level, text, page, idx in entries if page is None]
    if content_indices:
        start, end = min(content_indices), max(content_indices)
        span = max(1, end - start)
        for n, (level, text, page, idx) in enumerate(entries):
            if page is None:
                page = str(max(1, round((idx - start) / span * 26) + 1))
                entries[n] = (level, text, page, idx)
    return entries


def insert_before(anchor, element):
    anchor.addprevious(element)


def make_toc_entry(doc: Document, level: int, text: str, page: str):
    style_name = f"toc {level}"
    p = doc.add_paragraph(style=style_name if style_name in [s.name for s in doc.styles] else None)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.tab_stops.clear_all()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{text}\t{page}")
    set_run_font(r)
    if level == 2:
        r.bold = True
    return p


def rebuild_toc(doc: Document):
    toc_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "目  录" or p.text.strip() == "目 录")
    intro_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "绪  论" or p.text.strip() == "绪 论")
    anchor = doc.paragraphs[intro_idx]._element
    entries = collect_entries(doc, toc_idx)

    for p in doc.paragraphs[toc_idx + 1:intro_idx]:
        p._element.getparent().remove(p._element)

    begin_p = doc.add_paragraph()
    begin_p._p.append(fld_char("begin"))
    begin_p._p.append(instr_text(' TOC \\o "1-4" \\h \\z \\u '))
    begin_p._p.append(fld_char("separate"))
    insert_before(anchor, begin_p._element)

    for level, text, page, _ in entries:
        p = make_toc_entry(doc, level, text, page)
        insert_before(anchor, p._element)

    end_p = doc.add_paragraph()
    end_p._p.append(fld_char("end"))
    insert_before(anchor, end_p._element)


def remove_stale_toc_fields(doc: Document):
    # 清除WPS原先空目录留下的1-3级目录域，避免打开后继续按旧规则更新。
    for p in list(doc.element.findall(".//" + qn("w:p"))):
        instrs = p.findall(".//" + qn("w:instrText"))
        joined = "".join(i.text or "" for i in instrs)
        if 'TOC \\o "1-3"' in joined or 'TOC \\\\o "1-3"' in joined:
            parent = p.getparent()
            if parent is not None:
                parent.remove(p)


def set_pgnum(sect_pr, fmt=None, start=None):
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    if fmt is not None:
        pg.set(qn("w:fmt"), fmt)
    elif qn("w:fmt") in pg.attrib:
        del pg.attrib[qn("w:fmt")]
    if start is not None:
        pg.set(qn("w:start"), str(start))
    elif qn("w:start") in pg.attrib:
        del pg.attrib[qn("w:start")]


def clear_pgnum(sect_pr):
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is not None:
        sect_pr.remove(pg)


def ensure_page_number_sections(doc: Document):
    # 范文页码规则：摘要/Abstract等前置部分使用罗马数字；正文从“绪论”开始使用阿拉伯数字并重新从1开始。
    def compact(text):
        return "".join(text.split())

    abstract_idx = next(i for i, p in enumerate(doc.paragraphs) if compact(p.text) == "摘要")
    intro_idx = next(i for i, p in enumerate(doc.paragraphs) if compact(p.text) == "绪论")

    # 清理封面等摘要前页面的显式页码格式，避免封面被误设为罗马页码。
    for p in doc.paragraphs[:abstract_idx]:
        ppr = p._p.find(qn("w:pPr"))
        sect = ppr.find(qn("w:sectPr")) if ppr is not None else None
        if sect is not None:
            clear_pgnum(sect)

    # 找到“摘 要”之后遇到的第一个分节属性，该节包含中文摘要，设置为罗马页码I开始。
    for p in doc.paragraphs[abstract_idx:intro_idx]:
        ppr = p._p.find(qn("w:pPr"))
        sect = ppr.find(qn("w:sectPr")) if ppr is not None else None
        if sect is not None:
            set_pgnum(sect, "upperRoman", 1)
            break

    # “绪论”前一段作为前置部分结束位置；若没有分节符，则插入一个。
    prev = doc.paragraphs[intro_idx - 1]
    ppr = prev._p.get_or_add_pPr()
    front_sect = ppr.find(qn("w:sectPr"))
    if front_sect is None:
        base = doc.element.body.sectPr
        front_sect = deepcopy(base)
        ppr.append(front_sect)
    set_pgnum(front_sect, "upperRoman", None)

    # 文档最后的sectPr控制正文所在节，设置为阿拉伯数字并从1开始。
    body_sect = doc.element.body.sectPr
    set_pgnum(body_sect, "decimal", 1)


def main():
    doc = Document(DOCX)
    ensure_toc_styles(doc)
    set_heading_outline_levels(doc)
    add_update_fields_setting(doc)
    rebuild_toc(doc)
    remove_stale_toc_fields(doc)
    ensure_page_number_sections(doc)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
