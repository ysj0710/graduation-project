from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/Users/enjoy0710/Desktop/学校相关/毕业设计/graduation-project")
SRC = ROOT / "论文完善工作稿.docx"
OUT = ROOT / "基于Node_js_Vue的个人财务记账与可视化分析系统设计与实现-完善稿.docx"
ASSET_DIR = ROOT / "论文图表补充" / "论文完善版"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
SHOT_DIR = ROOT / "论文图表补充" / "功能截图"
FONT_PATH = "/System/Library/Fonts/Supplemental/Songti.ttc"
HEI_FONT_PATH = "/System/Library/Fonts/STHeiti Medium.ttc"


def f(size: int, path: str = FONT_PATH):
    return ImageFont.truetype(path, size=size, index=0)


F12, F14, F16, F18, F20, F22, F24 = [f(x) for x in (12, 14, 16, 18, 20, 22, 24)]
H16, H18, H20 = [f(x, HEI_FONT_PATH) for x in (16, 18, 20)]


def text_center(d, box, text, font=F16):
    x1, y1, x2, y2 = box
    lines = str(text).split("\n")
    dims = [d.textbbox((0, 0), line, font=font) for line in lines]
    heights = [b[3] - b[1] for b in dims]
    widths = [b[2] - b[0] for b in dims]
    total_h = sum(heights) + (len(lines) - 1) * 5
    y = y1 + (y2 - y1 - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        d.text((x1 + (x2 - x1 - w) / 2, y), line, fill=(0, 0, 0), font=font)
        y += h + 5


def vertical_text_center(d, box, text, font=F14, gap=1):
    x1, y1, x2, y2 = box
    chars = list(str(text))
    dims = [d.textbbox((0, 0), ch, font=font) for ch in chars]
    heights = [b[3] - b[1] for b in dims]
    widths = [b[2] - b[0] for b in dims]
    total_h = sum(heights) + (len(chars) - 1) * gap
    y = y1 + (y2 - y1 - total_h) / 2
    for ch, w, h in zip(chars, widths, heights):
        d.text((x1 + (x2 - x1 - w) / 2, y), ch, fill=(0, 0, 0), font=font)
        y += h + gap


def arrow_head(d, end, start, size=8):
    ex, ey = end
    sx, sy = start
    ang = math.atan2(ey - sy, ex - sx)
    pts = [
        (ex, ey),
        (ex - size * math.cos(ang - math.pi / 6), ey - size * math.sin(ang - math.pi / 6)),
        (ex - size * math.cos(ang + math.pi / 6), ey - size * math.sin(ang + math.pi / 6)),
    ]
    d.polygon(pts, fill="black")


def line_arrow(d, start, end, width=1, size=8):
    d.line([start, end], fill="black", width=width)
    arrow_head(d, end, start, size=size)


def rect(d, box, label="", font=F16, width=1):
    d.rectangle(box, fill="white", outline="black", width=width)
    if label:
        text_center(d, box, label, font)


def ellipse(d, center, label, size=(120, 48), font=F14):
    x, y = center
    w, h = size
    box = (int(x - w / 2), int(y - h / 2), int(x + w / 2), int(y + h / 2))
    d.ellipse(box, fill="white", outline="black", width=1)
    text_center(d, box, label, font)
    return box


def diamond(d, center, label, size=(112, 72), font=F16):
    x, y = center
    w, h = size
    pts = [(x, y - h // 2), (x + w // 2, y), (x, y + h // 2), (x - w // 2, y)]
    d.polygon(pts, fill="white", outline="black")
    d.line(pts + [pts[0]], fill="black", width=1)
    text_center(d, (x - w // 2, y - h // 2, x + w // 2, y + h // 2), label, font)


def point_on_rect(center, target, size):
    cx, cy = center
    tx, ty = target
    hw, hh = size[0] / 2, size[1] / 2
    dx, dy = tx - cx, ty - cy
    if dx == 0 and dy == 0:
        return int(cx), int(cy)
    scale = min(hw / abs(dx) if dx else 10**9, hh / abs(dy) if dy else 10**9)
    return int(cx + dx * scale), int(cy + dy * scale)


def point_on_diamond(center, target, size):
    cx, cy = center
    tx, ty = target
    hw, hh = size[0] / 2, size[1] / 2
    dx, dy = tx - cx, ty - cy
    if dx == 0 and dy == 0:
        return int(cx), int(cy)
    denom = abs(dx) / hw + abs(dy) / hh
    return int(cx + dx / denom), int(cy + dy / denom)


def draw_entity(d, center, name, attrs, positions, rect_size=(120, 58)):
    cx, cy = center
    rect_box = (cx - rect_size[0] // 2, cy - rect_size[1] // 2, cx + rect_size[0] // 2, cy + rect_size[1] // 2)
    rect(d, rect_box, name, F16)
    for label, pos in zip(attrs, positions):
        end = point_on_rect(center, pos, rect_size)
        d.line([end, pos], fill="black", width=1)
        width = max(62, min(112, len(label) * 14 + 22))
        ellipse(d, pos, label, size=(width, 34), font=F12)


def draw_relation(d, c1, c2, rel, label, a, b, s1, s2):
    dia = (72, 50)
    p1 = point_on_rect(c1, rel, s1)
    p2 = point_on_rect(c2, rel, s2)
    q1 = point_on_diamond(rel, p1, dia)
    q2 = point_on_diamond(rel, p2, dia)
    d.line([p1, q1], fill="black", width=1)
    d.line([q2, p2], fill="black", width=1)
    diamond(d, rel, label, dia, F14)
    d.text((p1[0] + (q1[0] - p1[0]) * 0.42, p1[1] + (q1[1] - p1[1]) * 0.42 - 18), a, fill="black", font=F14)
    d.text((p2[0] + (q2[0] - p2[0]) * 0.42, p2[1] + (q2[1] - p2[1]) * 0.42 - 18), b, fill="black", font=F14)


def save_architecture():
    img = Image.new("RGB", (1120, 720), "white")
    d = ImageDraw.Draw(img)
    rows = [
        ("表示层UI", ["Vue3", "Vue Router", "Pinia", "Element Plus", "ECharts"]),
        ("业务逻辑层BLL", ["用户认证", "记账管理", "账单导入", "统计分析", "预算预警", "消息通知", "风险监控", "系统管理"]),
        ("服务支撑层", ["Koa2", "RESTful API", "JWT鉴权", "Mongoose", "xlsx解析", "node-cron"]),
        ("数据持久层DAL", ["MongoDB", "finance_db", "集合索引", "文档模型"]),
        ("部署环境", ["Nginx", "HTTPS", "Node.js", "云服务器"]),
    ]
    x0, x1 = 35, 1085
    y = 32
    for title, items in rows:
        h = 112 if title == "业务逻辑层BLL" else 96
        rect(d, (x0, y, x1, y + h), "", F18)
        rect(d, (x0, y, x0 + 170, y + h), title, F20)
        area_x = x0 + 195
        cols = 4 if len(items) > 6 else len(items)
        gap = 18
        bw = int((x1 - area_x - gap * (cols - 1) - 20) / cols)
        bh = 38
        for i, item in enumerate(items):
            row = i // cols
            col = i % cols
            bx = area_x + col * (bw + gap)
            by = y + 20 + row * 48
            rect(d, (bx, by, bx + bw, by + bh), item, F16)
        y += h + 14
    path = ASSET_DIR / "图2-1 系统架构图.png"
    img.save(path)
    return path


def save_modules():
    img = Image.new("RGB", (1120, 520), "white")
    d = ImageDraw.Draw(img)
    root = (270, 25, 850, 84)
    user_box = (170, 170, 320, 230)
    admin_box = (800, 170, 950, 230)
    rect(d, root, "基于Node.js+Vue的个人财务记账与可视化分析系统", H20)
    rect(d, user_box, "用户端功能", H18)
    rect(d, admin_box, "管理端功能", H18)

    root_c = ((root[0] + root[2]) // 2, root[3])
    branch_y = 125
    user_c = ((user_box[0] + user_box[2]) // 2, user_box[1])
    admin_c = ((admin_box[0] + admin_box[2]) // 2, admin_box[1])
    d.line([root_c, (root_c[0], branch_y)], fill="black", width=1)
    d.line([(user_c[0], branch_y), (admin_c[0], branch_y)], fill="black", width=1)
    line_arrow(d, (user_c[0], branch_y), user_c, width=1, size=8)
    line_arrow(d, (admin_c[0], branch_y), admin_c, width=1, size=8)

    child_y1, child_y2 = 302, 485
    child_w = 42
    user_items = [
        "账号认证资料维护",
        "收支记账流水管理",
        "微信支付宝账单导入",
        "收支趋势可视化",
        "分类占比可视化",
        "消费结构可视化",
        "预算预警消息通知",
    ]
    admin_items = [
        "用户权限管理",
        "收支分类配置",
        "系统参数模板管理",
        "全局收支统计可视化",
        "分类消费分析可视化",
        "风险监控可视化",
        "数据导出日志审计",
    ]
    user_centers = [65 + i * 76 for i in range(len(user_items))]
    admin_centers = [600 + i * 76 for i in range(len(admin_items))]

    def draw_children(parent_box, centers, labels):
        parent_x = (parent_box[0] + parent_box[2]) // 2
        join_y = 270
        d.line([(parent_x, parent_box[3]), (parent_x, join_y)], fill="black", width=1)
        d.line([(centers[0], join_y), (centers[-1], join_y)], fill="black", width=1)
        for cx, label in zip(centers, labels):
            line_arrow(d, (cx, join_y), (cx, child_y1), width=1, size=6)
            box = (cx - child_w // 2, child_y1, cx + child_w // 2, child_y2)
            rect(d, box, "", F14)
            vertical_text_center(d, box, label, F14, gap=1)

    draw_children(user_box, user_centers, user_items)
    draw_children(admin_box, admin_centers, admin_items)
    path = ASSET_DIR / "图2-2 系统功能模块图.png"
    img.save(path)
    return path


USER_ATTRS = ["用户id", "用户名", "密码", "邮箱", "角色", "昵称", "头像", "创建时间", "更新时间"]
TX_ATTRS = ["记录id", "用户id", "收支类型", "金额", "分类名称", "备注", "交易日期", "创建时间", "更新时间"]
CAT_ATTRS = ["分类id", "分类名称", "图标id", "颜色", "收支类型", "默认标识", "排序值", "创建时间", "更新时间"]
ACC_ATTRS = ["账户id", "用户id", "账户名称", "账户类型", "图标", "颜色", "余额", "备注", "创建时间", "更新时间"]


def save_core_er():
    img = Image.new("RGB", (1100, 760), "white")
    d = ImageDraw.Draw(img)
    user, tx, cat = (295, 505), (550, 175), (805, 505)
    us, ts, cs = (72, 36), (86, 36), (72, 36)
    draw_entity(d, user, "用户", USER_ATTRS, [
        (210, 335), (125, 360), (70, 445), (75, 535), (135, 625),
        (260, 695), (385, 675), (470, 600), (455, 500)
    ], us)
    draw_entity(d, tx, "记账记录", TX_ATTRS, [
        (550, 45), (420, 58), (680, 58), (320, 125), (780, 125),
        (335, 225), (765, 225), (350, 300), (750, 300)
    ], ts)
    draw_entity(d, cat, "分类", CAT_ATTRS, [
        (890, 335), (975, 360), (1030, 445), (1025, 535), (965, 625),
        (835, 695), (715, 675), (625, 600), (650, 500)
    ], cs)
    draw_relation(d, user, tx, (420, 350), "记录", "1", "n", us, ts)
    draw_relation(d, tx, cat, (680, 350), "归属", "n", "1", ts, cs)
    path = ASSET_DIR / "图2-3 核心业务ER图.png"
    img.save(path)
    return path


def save_instance_er():
    img = Image.new("RGB", (740, 360), "white")
    d = ImageDraw.Draw(img)
    user, tx = (190, 190), (550, 190)
    us, ts = (72, 36), (86, 36)
    draw_entity(d, user, "用户", USER_ATTRS, [
        (190, 35), (95, 55), (45, 130), (45, 220), (90, 305),
        (180, 330), (285, 315), (340, 245), (330, 105)
    ], us)
    draw_entity(d, tx, "记账记录", TX_ATTRS, [
        (550, 35), (450, 55), (650, 55), (700, 130), (700, 220),
        (650, 305), (550, 330), (450, 305), (405, 130)
    ], ts)
    draw_relation(d, user, tx, (370, 190), "记录", "1", "n", us, ts)
    path = ASSET_DIR / "图2-4 用户记账实例ER图.png"
    img.save(path)
    return path


TABLES = [
    ("users表", [
        ("_id", "ObjectId", "PRIMARY KEY", "用户编号，MongoDB自动生成"),
        ("username", "String", "UNIQUE/REQUIRED", "登录用户名，3至20个字符"),
        ("password", "String", "REQUIRED", "加密后的登录密码"),
        ("email", "String", "UNIQUE/SPARSE", "用户邮箱，可用于验证码和找回密码"),
        ("role", "String", "ENUM", "用户角色，取值为admin或user"),
        ("nickname", "String", "DEFAULT", "用户昵称"),
        ("avatar", "String", "DEFAULT", "头像地址或图片数据"),
        ("createdAt", "Date", "DEFAULT", "创建时间"),
        ("updatedAt", "Date", "DEFAULT", "更新时间"),
    ]),
    ("transactions表", [
        ("_id", "ObjectId", "PRIMARY KEY", "记账记录编号"),
        ("userId", "ObjectId", "REF/INDEX", "所属用户，对应users._id"),
        ("type", "String", "ENUM/REQUIRED", "收支类型，income或expense"),
        ("amount", "Number", "REQUIRED", "交易金额，最小值为0"),
        ("category", "String", "REQUIRED", "分类名称，对应分类表中的name"),
        ("note", "String", "DEFAULT", "交易备注"),
        ("date", "Date", "DEFAULT/INDEX", "交易发生日期"),
        ("createdAt", "Date", "AUTO", "创建时间，timestamps生成"),
        ("updatedAt", "Date", "AUTO", "更新时间，timestamps生成"),
    ]),
    ("accounts表", [
        ("_id", "ObjectId", "PRIMARY KEY", "账户编号"),
        ("userId", "ObjectId", "REF/REQUIRED", "所属用户，对应users._id"),
        ("name", "String", "REQUIRED", "账户名称"),
        ("type", "String", "ENUM", "账户类型，cash、bank、alipay、wechat、other"),
        ("icon", "String", "DEFAULT", "账户图标"),
        ("color", "String", "DEFAULT", "账户颜色"),
        ("balance", "Number", "DEFAULT", "账户余额"),
        ("remark", "String", "DEFAULT", "备注"),
        ("createdAt", "Date", "DEFAULT", "创建时间"),
        ("updatedAt", "Date", "DEFAULT", "更新时间"),
    ]),
    ("categories表", [
        ("_id", "ObjectId", "PRIMARY KEY", "分类编号"),
        ("name", "String", "UNIQUE/REQUIRED", "分类名称"),
        ("iconId", "String", "DEFAULT", "分类图标编号"),
        ("color", "String", "DEFAULT", "分类颜色"),
        ("type", "String", "ENUM/REQUIRED", "分类类型，income或expense"),
        ("isDefault", "Boolean", "DEFAULT", "是否为系统默认分类"),
        ("sortOrder", "Number", "DEFAULT/INDEX", "分类排序值"),
        ("createdAt", "Date", "AUTO", "创建时间"),
        ("updatedAt", "Date", "AUTO", "更新时间"),
    ]),
    ("userconfigs表", [
        ("_id", "ObjectId", "PRIMARY KEY", "配置编号"),
        ("userId", "ObjectId", "REF/UNIQUE", "所属用户，对应users._id"),
        ("budget.monthly", "Number", "DEFAULT", "月度预算"),
        ("budget.yearly", "Number", "DEFAULT", "年度预算"),
        ("budget.alertThreshold", "Number", "DEFAULT", "预算预警阈值"),
        ("theme.background", "String", "DEFAULT", "主题背景"),
        ("theme.primaryColor", "String", "DEFAULT", "主题主色"),
        ("theme.glassBlur", "Number", "DEFAULT", "毛玻璃模糊度"),
        ("theme.pattern", "String", "DEFAULT", "背景图案"),
        ("theme.presetId", "String", "DEFAULT", "主题预设编号"),
        ("theme.customBgUrl", "String", "DEFAULT", "自定义背景图片地址"),
        ("notification.budgetAlert", "Boolean", "DEFAULT", "预算预警开关"),
        ("notification.riskAlert", "Boolean", "DEFAULT", "风险预警开关"),
        ("notification.systemAlert", "Boolean", "DEFAULT", "系统消息开关"),
        ("financialHealth.score", "Number", "DEFAULT", "财务健康评分"),
        ("financialHealth.riskLevel", "String", "ENUM/INDEX", "风险等级，low、medium、high"),
        ("financialHealth.lastCalculatedAt", "Date", "DEFAULT", "最近评分计算时间"),
        ("createdAt", "Date", "DEFAULT", "创建时间"),
        ("updatedAt", "Date", "DEFAULT", "更新时间"),
    ]),
    ("notifications表", [
        ("_id", "ObjectId", "PRIMARY KEY", "通知编号"),
        ("userId", "ObjectId", "REF/INDEX", "接收用户，对应users._id"),
        ("senderId", "ObjectId", "REF", "发送用户，可为空"),
        ("type", "String", "ENUM", "通知类型"),
        ("title", "String", "REQUIRED", "通知标题"),
        ("content", "String", "REQUIRED", "通知内容"),
        ("scope", "String", "ENUM", "发送范围，all或selected"),
        ("targetUserIds", "Array", "REF", "指定接收用户列表"),
        ("isRead", "Boolean", "DEFAULT/INDEX", "是否已读"),
        ("data", "Mixed", "DEFAULT", "扩展数据"),
        ("createdAt", "Date", "AUTO/INDEX", "创建时间"),
        ("updatedAt", "Date", "AUTO", "更新时间"),
    ]),
    ("operationlogs表", [
        ("_id", "ObjectId", "PRIMARY KEY", "日志编号"),
        ("userId", "ObjectId", "REF/REQUIRED", "操作用户，对应users._id"),
        ("action", "String", "REQUIRED", "操作类型"),
        ("details", "String", "DEFAULT", "操作详情"),
        ("createdAt", "Date", "AUTO", "创建时间"),
        ("updatedAt", "Date", "AUTO", "更新时间"),
    ]),
    ("systemsettings表", [
        ("_id", "ObjectId", "PRIMARY KEY", "系统设置编号"),
        ("systemName", "String", "DEFAULT", "系统名称"),
        ("currency", "String", "DEFAULT", "默认币种"),
        ("defaultBudget", "Number", "DEFAULT", "默认预算"),
        ("createdAt", "Date", "AUTO", "创建时间"),
        ("updatedAt", "Date", "AUTO", "更新时间"),
    ]),
    ("messagetemplates表", [
        ("_id", "ObjectId", "PRIMARY KEY", "模板编号"),
        ("low", "String", "DEFAULT", "低风险提示模板"),
        ("medium", "String", "DEFAULT", "中风险提示模板"),
        ("high", "String", "DEFAULT", "高风险提示模板"),
        ("createdAt", "Date", "AUTO", "创建时间"),
        ("updatedAt", "Date", "AUTO", "更新时间"),
    ]),
    ("verificationcodes表", [
        ("_id", "ObjectId", "PRIMARY KEY", "验证码编号"),
        ("email", "String", "REQUIRED", "接收验证码的邮箱"),
        ("code", "String", "REQUIRED", "验证码内容"),
        ("type", "String", "ENUM", "验证码类型，register或resetPassword"),
        ("expiresAt", "Date", "TTL INDEX", "过期时间，过期后自动删除"),
        ("createdAt", "Date", "DEFAULT", "创建时间"),
    ]),
]


def set_font(run, name="宋体", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def style_paragraph(p, first_indent=True):
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(0)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    for run in p.runs:
        set_font(run)


def caption(doc, anchor, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    for run in p.runs:
        set_font(run)
    anchor.addprevious(p._element)
    return p


def para(doc, anchor, text, style="Body Text", first_indent=True):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    style_paragraph(p, first_indent=first_indent and not style.startswith("Heading"))
    anchor.addprevious(p._element)
    return p


def pic(doc, anchor, path, cap, width=15.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    anchor.addprevious(p._element)
    caption(doc, anchor, cap)


def cell_shade(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_width(cell, cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    w = OxmlElement("w:tcW")
    w.set(qn("w:w"), str(int(cm * 567)))
    w.set(qn("w:type"), "dxa")
    tc_pr.append(w)


def table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_table(doc, anchor, cap, rows):
    table_no = cap.split(" ", 1)[0]
    para(doc, anchor, f"{cap.split(' ', 1)[1]}详细设计如{table_no}所示。")
    caption(doc, anchor, cap)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        table_borders(table)
    widths = [3.6, 3.0, 3.0, 7.0]
    headers = ["字段名", "数据类型", "字段属性", "字段描述"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell_width(cell, widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cell_width(cells[i], widths[i])
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_font(r, size=9)
    anchor.addprevious(table._element)
    para(doc, anchor, "", first_indent=False)


def test_table(doc, anchor, cap, rows):
    caption(doc, anchor, cap)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        table_borders(table)
    widths = [2.2, 3.0, 5.0, 4.2, 2.2]
    headers = ["编号", "测试模块", "测试内容", "预期结果", "结果"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell_width(cell, widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cell_width(cells[i], widths[i])
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci in (0, 1, 4) else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_font(r, size=9)
    anchor.addprevious(table._element)
    para(doc, anchor, "", first_indent=False)


EXPANSION = {
    "requirements": [
        "从业务流程角度看，普通用户的核心路径可以概括为注册登录、进入首页、创建或导入记账记录、查看统计分析、根据分析结果调整预算设置。该路径覆盖了系统最主要的使用场景，也决定了系统设计时必须优先保证记账入口清晰、列表查询稳定、统计结果准确。若用户需要逐条记录日常消费，系统应当支持快速录入金额、分类、日期和备注；若用户已经在微信或支付宝中积累了大量交易数据，系统则应通过批量导入降低迁移成本。",
        "管理员的需求与普通用户不同，重点不在单个用户的记账过程，而在系统整体运行情况和数据治理。管理员需要查看用户数量、交易规模、风险用户分布等运营指标，也需要维护收入和支出分类，保证分类体系稳定统一。对于预算风险较高的用户，管理员端需要能够查看风险等级和风险原因，并通过系统通知发送提醒。这些需求在项目中主要由admin路由和管理员端视图共同实现。",
        "系统的数据分析需求不是简单地展示交易列表，而是要将流水数据转换为可理解的财务信息。用户关心的是本月收入、支出、结余、分类占比和消费趋势，管理员关心的是整体交易规模、活跃情况和风险分布。因此后端必须提供聚合统计接口，前端必须提供图表展示能力。项目中使用MongoDB聚合管道完成按类型、分类和日期的统计，再由ECharts将结果渲染为趋势图、饼图和统计卡片。",
        "安全需求贯穿整个系统。个人财务数据具有隐私性，任何用户都只能访问自己的记账记录、账户配置和通知信息。系统通过JWT在接口层面识别用户身份，交易、账户、通知等查询均以用户id作为过滤条件。管理员虽然可以查看系统层面的统计和用户列表，但后台接口仍需要校验角色，避免普通用户绕过前端路由直接访问管理接口。",
    ],
    "feasibility": [
        "从项目规模看，本系统功能覆盖用户端和管理员端，但各模块边界较清晰，适合采用前后端分离方式逐步实现。前端负责页面组织和交互状态，后端负责数据校验和持久化，数据库负责保存文档数据。这样的职责划分能够降低模块之间的耦合度，使记账、导入、通知、风险评估等功能可以独立开发和测试。",
        "从开发环境看，前端Vite提供快速启动和热更新能力，能够提高页面调试效率；后端Node.js使用JavaScript语言，与前端语言一致，减少了开发过程中的语言切换成本。MongoDB的文档结构与JavaScript对象天然接近，Mongoose又提供Schema约束和模型方法，使数据库操作既保持灵活，又具备一定规范性。",
        "从维护成本看，本系统不依赖商业收费组件，主要技术均可在普通开发环境中部署运行。前端打包后生成静态资源，可通过Nginx部署；后端作为Node.js服务运行，数据库使用MongoDB保存数据。对于毕业设计和中小规模应用而言，这种部署方式成本较低，后续维护主要集中在功能迭代和数据备份上。",
        "从用户学习成本看，系统页面以财务管理常见概念组织，包括交易记录、统计分析、消费分析、个性设置、消息中心等。用户不需要理解数据库或接口细节，只需要按照页面提示完成操作。管理员端也采用常见的后台管理布局，列表、筛选、表单、弹窗和分页等交互方式符合一般后台系统使用习惯。",
    ],
    "related": [
        "Vue3在本系统中的作用不仅是渲染页面，还承担了组件复用和状态响应的基础能力。首页统计卡片、交易筛选表单、消息列表、管理员数据表格等界面都可以拆分为相对独立的组件。组件化设计使页面结构更清晰，也便于后续对某个功能局部调整，例如只修改分类图标组件而不影响交易列表的整体逻辑。",
        "Vue Router负责组织用户端和管理员端页面。本项目路由文件中定义了/dashboard、/transactions、/statistics、/consumption、/personal、/notifications等用户端路径，也定义了/admin/dashboard、/admin/users、/admin/risk、/admin/finance-stats、/admin/settings-category等管理端路径。路由守卫通过本地保存的Token和用户角色判断访问权限，是前端权限控制的重要组成部分。",
        "Koa2采用中间件机制处理请求，适合本系统这种接口型后端。请求进入系统后，先经过跨域和请求体解析，再进入业务路由；需要登录的接口继续经过JWT校验，校验通过后再执行具体业务。相较于把所有逻辑写在一个入口文件中，路由拆分让后端结构更清晰，也便于按模块排查错误。",
        "MongoDB适合保存本系统的配置类和通知类数据。用户配置中包含预算、主题、通知开关和财务健康评分等嵌套结构，如果使用关系型数据库，需要拆分多个表并维护关联；在MongoDB中可以通过一个userconfigs文档保存用户的完整配置。对于毕业设计项目来说，这种文档模型能够减少复杂关联，同时保留后续扩展字段的空间。",
    ],
    "architecture": [
        "表示层主要由Vue3单页面应用构成，负责向用户展示系统功能。用户在浏览器中完成登录、记账、查看图表和设置预算等操作，页面上的按钮、表单、表格和图表都属于表示层内容。表示层不直接访问数据库，而是通过Axios请求后端接口，从而保证前端只关注交互和展示。",
        "业务逻辑层由Koa路由、中间件和业务处理函数组成。该层负责判断请求是否合法、参数是否完整、分类与收支类型是否匹配、预算是否达到预警阈值、管理员是否具有访问权限等。业务逻辑层是系统中最关键的部分，它连接前端请求和数据库操作，保证系统行为符合需求分析中的约束。",
        "数据持久层由MongoDB和Mongoose模型组成。Mongoose模型定义了字段类型、默认值、枚举范围和索引信息，例如Transaction模型限定type只能为income或expense，amount必须为非负数，Category模型限定分类类型，UserConfig模型限定风险等级。通过模型约束，系统可以在写入数据库前过滤一部分非法数据。",
        "部署层面，前端资源可以由Nginx托管，后端Node.js服务负责处理/api请求，数据库部署在同一服务器或独立数据库服务中。由于前后端分离，后续如果需要扩展为多端应用，只要接口保持稳定，移动端或小程序端也可以复用后端服务。这种架构与绪论中提出的多端使用和后续可扩展目标保持一致。",
    ],
    "modules": [
        "用户信息模块是系统入口模块，承担身份确认和个人资料维护任务。注册时用户需要输入用户名、邮箱、密码和验证码，后端验证验证码有效后创建用户；登录时后端根据用户名查询用户并使用bcryptjs校验密码，成功后返回Token。个人资料维护包括昵称、用户名、头像和密码等信息，相关接口集中在auth和user路由中。",
        "记账管理模块是系统最核心的业务模块。用户可以新增单条记录，也可以通过items数组一次提交多条记录。后端在保存前会校验分类与收支类型是否匹配，避免把收入分类用于支出记录。交易列表查询支持分页和条件筛选，能够满足用户按时间、分类、类型和备注查找流水的需求。",
        "统计分析模块负责把原始交易记录转化为汇总信息。系统按时间范围统计收入总额、支出总额和结余，按分类统计各类支出的金额和占比，按日期统计每日收支趋势。前端将这些数据组织成统计卡片、折线图、柱状图和饼图，让用户能够从整体上理解自己的消费结构。",
        "预算通知模块结合了用户配置、定时任务和通知集合。用户在个性设置中配置月度预算和预警阈值，后端在新增支出或定时任务触发时计算当月支出总额。当预算使用率达到阈值时，系统写入一条预算预警通知，并在消息中心显示。该模块体现了系统从被动记账向主动提醒的扩展。",
        "后台管理模块主要服务管理员，包含管理总览、用户管理、风险监控、分类管理、系统设置、消息模板、数据导出和操作日志等页面。管理员可以从系统层面查看用户数量和交易规模，也可以维护分类和模板。操作日志用于记录关键后台操作，为后续审计和问题追踪提供依据。",
    ],
    "database": [
        "数据库设计遵循“用户隔离、记录独立、配置集中、通知可追踪”的原则。用户相关的交易、账户、配置和通知都通过userId关联到users集合，这样既能保证普通用户只能访问自己的数据，也便于管理员在后台按用户维度进行统计分析。交易记录本身不直接嵌入用户文档，而是单独保存在transactions集合中，避免用户文档随交易数量增长而过大。",
        "transactions集合是统计分析的基础，因此项目为其设置了多个索引，包括userId和date的组合索引、userId和type的组合索引、date索引、type和date组合索引以及category索引。这些索引分别服务于用户流水查询、类型筛选、管理员按日期统计、收支聚合和分类统计等场景。索引设计与系统查询方式相匹配，能够减少全表扫描的概率。",
        "categories集合用于维护收入和支出分类。分类字段不直接保存为ObjectId，而是在Transaction中以category字符串保存分类名称，这种设计降低了记账记录展示时的关联查询成本，也符合个人记账系统中分类名称稳定、展示频繁的特点。分类表中仍保留type、iconId、color、isDefault和sortOrder等字段，用于前端区分收入支出、展示图标颜色和排序。",
        "userconfigs集合保存预算、主题、通知开关和风险评分等用户个性化信息。与把这些字段分散到用户表不同，单独的配置集合可以避免users集合承担过多职责。通知集合notifications则保存预算预警、系统通知、风险提醒和交易相关通知，包含isRead字段和createdAt索引，便于消息中心查询未读消息和按时间倒序展示。",
        "论文表结构中没有列出__v字段，是因为它并非业务设计字段。Mongoose默认会为文档生成versionKey，字段名通常为__v，用于内部版本控制。在本项目代码中，无论用户认证、记账统计、风险计算还是后台管理，都没有读取或写入__v。因此在毕业论文的数据库设计中将其排除，更符合“展示业务字段和必要系统字段”的写作规范。",
        "字段命名方面，项目代码采用JavaScript中常见的驼峰命名和点路径表示嵌套字段。例如alertThreshold表示预算预警阈值，financialHealth.riskLevel表示财务健康风险等级，theme.primaryColor表示主题主色。论文表格中保留这些真实字段名，能够使数据库设计与源码保持一致，也便于答辩时根据字段追溯到具体模型。",
        "数据完整性方面，系统主要通过Mongoose Schema、枚举约束、必填字段和业务校验共同保证。例如User模型限制用户名长度，Transaction模型限制type枚举和amount最小值，Category模型限制分类类型，UserConfig模型限制风险等级枚举。对于MongoDB这种文档数据库而言，合理的Schema约束可以弥补无固定表结构带来的随意性。",
    ],
    "backend": [
        "后端入口文件将数据库连接和路由注册集中处理，保证服务启动后各模块具备统一访问入口。MongoDB连接成功后会启动定时任务，这说明预算预警并不完全依赖用户主动操作，也可以由后台定期扫描触发。这样的设计使系统即使在用户没有新增记录时，也能定期检查预算状态和月度统计情况。",
        "auth路由承担身份相关逻辑，接口包括发送验证码、验证验证码、注册、登录、Token校验、获取当前用户、忘记密码、注销账号和修改密码等。注册接口要求邮箱验证码，登录接口根据用户名和密码生成Token，Token中包含用户id、用户名、邮箱和角色。通过这些接口，系统形成了比较完整的认证闭环。",
        "transactions路由是用户端交易功能的核心。GET接口支持分页、类型、分类、日期和关键词筛选；POST接口既支持单条记录，也支持批量items提交；PUT接口用于更新交易记录；DELETE接口用于删除记录；statistics、month-stats和daily-stats接口用于不同粒度的统计展示。接口设计围绕前端页面实际需要展开，没有把统计计算全部交给前端。",
        "admin路由覆盖后台大部分功能，接口数量较多，包括dashboard-stats、trend-data、risk-users、risk-distribution、users、finance-stats、category-analysis、system-settings、categories、message-templates、operation-logs和export等。管理员端功能复杂，因此将后台接口集中到admin路由中，便于统一加入管理员权限校验。",
        "notifications路由负责用户消息，支持获取通知列表、获取未读数量、标记单条已读、全部已读、删除和批量删除，也支持预算预警检查。通知功能在系统中起到连接预算、风险和用户提醒的作用，使分析结果不只停留在图表展示，而能够通过消息主动反馈给用户。",
        "后端接口的返回格式以JSON为主，通常包含message字段和具体数据字段。对于列表类接口，系统会返回数据数组、总数、当前页和总页数；对于新增、修改和删除类接口，系统会返回操作结果和必要的对象信息。统一的响应风格能够降低前端处理复杂度，前端只需要根据状态码和message提示用户即可。",
        "异常处理方面，项目中的路由大多使用try-catch包裹数据库操作。当数据库查询失败、参数不合法或业务校验不通过时，后端会返回对应状态码和错误信息。虽然毕业设计阶段的错误处理还可以继续细化，但当前实现已经能够覆盖常见接口异常，避免后端错误直接暴露为无响应或服务崩溃。",
    ],
    "auth": [
        "用户注册流程首先校验请求体中的username、password、email和code是否完整，再根据email和code查询verificationcodes集合。验证码不存在或过期时，后端返回错误信息，不创建用户。验证码通过后，系统检查用户名是否已存在，然后使用bcrypt.hash对密码加密，最后创建User文档并删除已使用的验证码。该流程能够避免无效邮箱注册和明文密码存储。",
        "登录流程相对直接，但同样包含必要的安全处理。后端先根据username查询用户，如果用户不存在则返回统一的错误提示，避免暴露账号是否存在；如果用户存在，则使用bcrypt.compare比较明文密码和数据库中的哈希密码。密码正确后生成有效期为24小时的JWT，并将用户基础信息返回给前端。",
        "前端路由守卫根据localStorage中的Token和user信息判断访问权限。管理员用户访问普通用户路径时会被重定向到/admin，普通用户访问/admin路径时会被重定向到/dashboard。这种前端控制提升了使用体验，但真正的安全边界仍在后端管理员接口的角色校验，因为用户可以绕过前端直接发送HTTP请求。",
        "密码修改、找回密码和账号注销属于敏感操作。项目中相关接口都要求用户身份或验证码验证，避免未登录用户直接修改账号数据。对于毕业设计系统而言，这些处理已经覆盖了常见账号安全场景。后续如果投入真实生产环境，还可以增加登录失败次数限制、刷新Token机制和更细粒度的审计日志。",
    ],
    "transactions": [
        "新增记账时，后端会调用分类类型校验函数validateCategoryType。该函数根据分类名称查询Category集合，如果分类存在，就判断分类type是否与交易type一致；如果分类不存在，则不阻止写入。这种处理兼顾了系统预设分类的规范性和用户导入账单时分类名称可能不完全匹配的灵活性。",
        "批量记账使用items数组提交，每个条目可以包含自己的category、amount、note和type。后端逐条过滤缺少金额或分类的条目，并跳过分类类型不匹配的数据。只有待保存数组不为空时才执行insertMany写入数据库。这种方式比前端循环逐条调用接口更高效，也能减少网络请求次数。",
        "交易查询接口使用page和pageSize控制分页，通过skip和limit返回指定页数据，并返回total、page、pageSize和totalPages。分页信息对前端表格很重要，用户可以清楚知道当前查询结果总量，也能避免一次性加载大量历史记录造成页面卡顿。日期范围筛选和备注关键词搜索则满足了用户查找特定流水的实际需求。",
        "统计接口使用聚合管道而不是普通查询后在前端计算。收入、支出、分类占比和每日趋势都依赖MongoDB的$match、$group和$sort等操作完成。这样做可以保证统计逻辑集中在后端，减少前端重复计算，并且便于管理员端和用户端复用类似的统计口径。",
        "删除交易使用findOneAndDelete并同时匹配_id和userId，保证用户不能删除不属于自己的记录。更新交易时也通过_id和userId共同查询，找到记录后再按请求体更新字段。这种写法是用户数据隔离的重要体现，也是个人财务系统必须满足的基本安全要求。",
        "从用户体验角度看，记账模块需要在“快速录入”和“信息完整”之间取得平衡。如果表单字段过多，用户会觉得记账成本高；如果字段过少，后续统计又缺乏维度。本系统将交易记录控制在收支类型、金额、分类、备注和日期等核心字段上，既能支撑统计分析，又不会让单次录入过于复杂。",
        "从数据质量角度看，分类字段对后续分析影响很大。如果用户长期使用不规范分类，分类占比图就会失去参考价值。因此系统在新增记录时校验分类与收支类型，管理员端又提供分类维护能力，两者共同保证收入分类和支出分类不会混用，增强统计结果的可信度。",
    ],
    "risk": [
        "账单导入模块在实现上需要处理不同平台的文件格式差异。微信账单中表头可能不在第一行，代码会在前30行中查找包含“交易时间”的表头行，再根据表头定位交易时间、交易类型、交易对方、商品、收支和金额等列。对于金额字段，系统会去除人民币符号和其他非数字字符，再转换为数值。",
        "支付宝账单与微信账单的列结构不同，导入接口单独提供/alipay路径处理。系统在解析后将外部账单转换为统一的Transaction结构，包括type、amount、category、note和date等字段。无论来源是微信还是支付宝，最终进入数据库的都是同一种交易文档，这有利于后续统计分析保持一致。",
        "预算预警逻辑既出现在新增支出之后，也可以由定时任务周期性执行。系统读取UserConfig中的monthly预算和alertThreshold阈值，统计当月expense类型交易总额，计算usageRate。如果当天已经发送过预算预警，则不重复创建通知，避免用户在短时间内收到大量相同提醒。",
        "风险评估算法的三个维度具有明确含义。预算使用率反映用户本月支出是否接近预算上限，收支平衡反映支出相对收入是否过高，交易频率反映用户消费行为是否过于频繁。系统采用40%、30%、30%的权重计算综合分，是因为预算使用率与个人财务风险关系最直接，因此权重最高。",
        "管理员风险监控并不是简单展示风险等级，而是将UserConfig中的financialHealth字段与用户信息、月度收支和预算数据结合起来。管理员可以根据风险等级筛选用户，也可以发送风险提醒通知。该设计让风险评估从算法结果延伸到管理操作，增强了系统的实用性。",
    ],
    "frontend": [
        "前端用户端采用布局组件承载侧边导航和内容区域，各功能页面作为子路由加载。Dashboard作为首页直接引入，其他页面使用懒加载函数按需加载。路由文件中还包含延迟预取逻辑，在首页加载后利用import.meta.glob预取用户端和管理员端页面，减少后续页面切换时的等待时间。",
        "交易记录页面是用户高频使用页面，因此交互设计重点在筛选效率和表单反馈。用户可以通过类型、分类、日期和关键词缩小查询范围；新增或编辑记录时，表单应保证金额、分类、日期等必填项完整；删除记录需要明确反馈结果，避免用户误以为操作失败。前端的这些交互与后端分页和校验逻辑共同保证了记账体验。",
        "统计分析页面依赖ECharts完成可视化展示。图表展示的价值不在于形式复杂，而在于帮助用户快速理解收支变化。收支趋势可以显示某段时间内收入和支出的变化方向，分类占比可以帮助用户识别主要消费领域，消费分析页面则进一步从结构和趋势角度分析用户消费习惯。",
        "个性设置页面承担用户配置维护，包括主题、预算和通知偏好。主题设置保存背景、主色、模糊度、图案、预设编号和自定义背景地址等字段；预算设置保存月度预算、年度预算和预警阈值；通知设置控制预算预警、风险预警和系统通知是否开启。这些配置最终保存在userconfigs集合中。",
        "消息中心页面展示预算预警、系统通知、风险提醒和交易相关消息。前端需要区分已读和未读状态，并提供标记已读、全部已读和删除等操作。未读数量接口可以用于导航栏或页面入口的提示，使用户及时看到预算和风险相关提醒。",
        "主题系统是本项目用户体验设计的一部分。系统不是只提供固定配色，而是将背景、主色、玻璃模糊度、图案和自定义背景地址等配置保存到用户配置中。这样用户重新登录后仍能恢复自己的界面偏好，也使系统在视觉体验上区别于传统单调的记账页面。",
        "前端页面还需要处理加载状态和空数据状态。交易列表为空时应给出明确提示，统计数据不足时图表不应出现异常，接口请求过程中需要有加载反馈，操作完成后需要提示成功或失败。这些细节虽然不属于核心算法，但直接影响用户是否愿意长期使用系统。",
    ],
    "admin": [
        "管理员端首页强调概览能力。系统通过dashboard-stats接口聚合用户总数、今日新增用户、活跃用户、月度交易笔数、月度交易金额和风险用户数量。管理员不需要进入每个功能页面就能了解系统整体状态，这符合后台管理系统先看总览、再处理异常的使用习惯。",
        "用户管理页面以列表为中心，支持分页、查看、创建、编辑、删除和批量操作。管理员创建用户时，后端会检查用户名唯一性并创建对应用户文档；更新用户时可以修改用户名、邮箱、角色、预算和主题等信息；删除用户时会移除用户文档。由于用户删除属于高风险操作，实际使用中应配合二次确认。",
        "分类管理页面维护系统分类体系。分类包含name、iconId、color、type、isDefault和sortOrder字段，管理员可以新增、修改、删除和排序分类。分类设计会直接影响用户记账时的可选项和统计分析中的分类维度，因此分类管理需要保证名称清晰、类型准确、排序合理。",
        "系统设置和消息模板属于后台配置功能。系统设置保存系统名称、默认币种和默认预算，影响前端显示和用户默认配置；消息模板保存低、中、高风险提醒文案，风险监控发送提醒时可以使用这些模板。将这些配置保存到数据库中，可以避免每次调整都修改代码。",
        "操作日志模块用于记录管理员关键行为，字段包括userId、action、details、createdAt和updatedAt。虽然毕业设计阶段的日志内容相对简单，但它体现了后台系统应具备的可追溯性。后续如果扩展，可记录请求IP、操作对象、修改前后内容等信息，进一步增强审计能力。",
        "数据导出功能服务于管理员的数据备份和分析需求。admin路由中的export接口可以根据不同类型导出用户、交易或统计数据，并支持JSON、CSV、Excel等格式。对于个人财务系统而言，导出能力也具有数据可迁移意义，避免用户数据被锁定在单一系统中。",
        "管理员端的风险提醒功能与普通消息通知不同，它带有一定管理干预性质。管理员在风险监控中发现高风险或中风险用户后，可以通过单独提醒或批量提醒发送消息。这样系统不仅能自动识别风险，还能为人工管理提供操作入口，形成“识别、查看、提醒、反馈”的闭环。",
    ],
    "testing": [
        "测试工作从功能正确性、数据隔离性、异常处理和页面表现四个方面展开。功能正确性关注用户能否完成登录、记账、查询、统计、导入和通知操作；数据隔离性关注用户只能操作自己的数据；异常处理关注缺少参数、错误密码、分类不匹配、无效Token等情况；页面表现关注图表、表格和表单是否能够正常显示。",
        "项目中已有Jest和Supertest相关测试文件，包括auth.test.js、transaction.test.js和risk.test.js。认证测试覆盖注册、重复用户名、缺少字段、登录成功、错误密码和不存在用户等场景；交易测试覆盖添加支出、添加收入、缺少必要字段、列表查询、类型筛选、统计和删除；风险测试覆盖预算使用率、收支平衡、交易频率和综合评分。",
        "由于系统使用MongoDB，测试环境与正式数据库需要隔离。项目测试文件使用finance_test作为测试数据库，并在测试结束后清理数据。这样可以避免测试过程中插入的用户和交易记录污染正式finance_db数据库，也能保证每次测试运行时数据状态相对可控。",
        "功能测试用例采用黑盒测试思路，从用户操作结果判断功能是否符合预期。例如输入正确用户名和密码应成功登录，输入错误密码应提示错误；新增支出记录后交易列表应出现该记录，预算达到阈值后消息中心应出现提醒；普通用户访问管理员路径应被拦截或重定向。这些测试直接对应用户可感知的系统行为。",
        "性能测试在毕业设计阶段主要关注接口响应和页面可用性。交易记录采用分页查询，统计接口使用数据库聚合，前端页面采用懒加载和延迟预取，这些设计都服务于性能目标。虽然未进行大规模压测，但从实现方式看，系统已经避免了明显的一次性加载全部数据和前端重复统计问题。",
        "边界测试主要关注输入为空、金额为负数、分类类型不匹配、验证码过期、Token无效、删除不存在记录等情况。这些场景在真实使用中经常出现，如果处理不当会导致数据错误或用户困惑。项目中的后端路由针对许多异常情况返回了明确提示，为前端展示错误信息提供了基础。",
        "安全测试主要关注越权访问。普通用户不能访问管理员接口，也不能通过修改请求参数查询或删除其他用户的交易记录。系统在交易更新和删除时同时匹配记录id和用户id，在管理员路由中校验角色，这些设计能够防止最常见的水平越权和垂直越权问题。",
    ],
    "test_analysis": [
        "认证模块测试说明系统能够正确区分合法用户和非法用户。注册流程对必要字段、验证码和用户名唯一性进行检查，登录流程对用户名和密码进行校验，并在成功后返回Token。该模块是系统安全的第一道入口，如果认证模块存在缺陷，后续记账和管理功能都无法保证数据安全。",
        "记账模块测试说明交易记录的基本增删查统功能能够运行。添加记录后能够在列表中查询到，按类型筛选能够返回对应数据，统计接口能够正确计算收入、支出和结余。由于统计图表依赖这些接口，后端统计结果正确是前端可视化可信的前提。",
        "风险算法测试覆盖了多个边界区间，例如预算使用率超过100%、90%至100%、80%至90%、60%至80%以及低于60%的情况，也覆盖了支出大于收入、支出接近收入和交易频率过高等情况。通过这些测试可以确认风险等级划分与代码中的规则一致。",
        "系统仍有可改进之处。第一，导入功能可以增加更多银行账单格式支持；第二，账号注销和用户删除属于高风险操作，后续可增加更严格的确认机制；第三，统计接口在数据量持续增加后可以增加缓存或预计算机制；第四，测试可以进一步覆盖前端组件和端到端流程，使页面交互也纳入自动化验证范围。",
        "测试结果也说明，系统的主要风险不在单个功能是否能够运行，而在复杂场景下多个模块之间的数据一致性。例如导入账单会影响交易记录和账户余额，新增支出会影响预算预警和统计图表，修改分类会影响后续记账选择和分类分析。因此后续测试应进一步增加跨模块场景，验证一处数据变化能否在相关页面中正确反映。",
        "从毕业设计完成度看，系统已经具备完整的前后端闭环、真实数据库模型、用户端和管理员端页面、自动化测试文件和可部署的技术栈。测试章节不是为了证明系统没有任何问题，而是说明系统主要功能达到了设计目标，并明确后续可以继续优化的方向。",
    ],
    "social": [
        "个人财务管理系统的社会价值不在于替代专业理财服务，而在于帮助普通用户建立基础的财务意识。许多用户虽然每天都在使用移动支付，但并不会主动回顾自己的消费结构。本系统通过分类统计和趋势展示，把原本零散的支付记录转化为可理解的消费信息，有助于用户发现长期被忽略的小额高频支出。",
        "预算预警功能可以在用户消费接近阈值时给予提醒，这种提醒对学生和刚参加工作的年轻群体尤其有意义。相比月底才发现超支，提前预警能够让用户及时调整后续消费安排。财务健康风险评估虽然不能替代专业财务建议，但可以作为自我检查工具，帮助用户关注预算、收支平衡和消费频率。",
        "从社会治理角度看，理性消费和个人财务规划有助于减少非理性借贷和过度消费带来的风险。系统如果进一步推广，可以作为高校学生财务教育、家庭预算管理或个人消费复盘的辅助工具。它通过低门槛的数据记录和图表反馈，让用户逐步形成量入为出的消费习惯。",
        "从环境角度看，电子化记账可以减少纸质账本和纸质票据整理的依赖。用户通过导入电子账单和在线查看统计结果即可完成大部分财务记录工作。系统采用Web方式访问，也减少了传统客户端分发、重复安装和频繁更新带来的资源消耗。",
        "从技术伦理角度看，财务数据属于高度敏感信息，系统设计必须坚持最小必要原则。毕业设计版本已经通过密码加密、Token鉴权和用户数据隔离降低风险，但如果未来面向真实用户开放，还需要补充隐私政策、数据导出与删除机制、备份加密和异常访问监控，确保系统的社会价值不以牺牲用户隐私为代价。",
        "从可持续使用角度看，个人财务管理工具的价值来自长期积累。用户记录一两天数据并不能形成有效分析，只有持续记录和定期复盘，图表和风险评估才会产生参考意义。因此系统在功能上提供导入、统计、预算、消息和导出能力，目的就是降低长期使用阻力，让用户愿意持续保存和查看自己的财务数据。",
        "从项目研究角度看，本系统虽然是毕业设计作品，但覆盖了真实Web系统开发中的多个关键问题，包括身份认证、权限控制、文档数据库设计、数据聚合、文件导入、定时任务、图表可视化和后台管理。通过这些模块的设计与实现，论文能够体现作者对软件工程实践流程的理解，而不仅是对单一技术点的简单使用。",
    ],
}


def add_expand(doc, anchor, key):
    for text in EXPANSION[key]:
        para(doc, anchor, text)


def delete_allowed_block(doc):
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "系统分析":
            start = i
        if t == "结论":
            end = i
            break
    if start is None or end is None:
        raise RuntimeError("未找到可替换章节边界")
    anchor = doc.paragraphs[end]._element
    for p in doc.paragraphs[start:end]:
        p._element.getparent().remove(p._element)
    return anchor


def build_content(doc, anchor, paths):
    para(doc, anchor, "1 系统分析", "Heading 2", False)
    para(doc, anchor, "1.1 需求分析", "Heading 3", False)
    para(doc, anchor, "本系统面向个人日常收支管理场景，目标是在浏览器端提供一个无广告、可持续使用、兼顾记账效率和数据分析深度的财务管理工具。结合摘要和绪论中提出的移动支付普及、个人消费数据沉淀、现有记账工具功能受限等问题，系统需要围绕“记录、分析、预警、管理”四类需求展开设计。")
    para(doc, anchor, "普通用户侧需要完成账号注册登录、个人资料维护、日常收支记录、按类型和分类查询流水、微信和支付宝账单导入、统计图表查看、预算设置、预算预警消息查看、主题个性化设置等操作。管理员侧需要完成用户管理、分类管理、系统基础设置、风险用户监控、消息模板维护、通知发送、操作日志查看和数据导出等管理工作。")
    para(doc, anchor, "在非功能需求方面，系统需要保证用户数据隔离和接口访问安全，所有需要登录的接口均通过JWT进行身份校验；前端页面需要具备较好的响应速度和交互反馈；后端接口需要支持分页、筛选和聚合统计，避免一次性返回过多数据；数据库设计需要能够支撑用户、记账记录、分类、账户、通知、配置等文档数据的长期扩展。")
    add_expand(doc, anchor, "requirements")
    para(doc, anchor, "1.2 可行性分析", "Heading 3", False)
    para(doc, anchor, "1.2.1 技术可行性", "Heading 4", False)
    para(doc, anchor, "本系统采用Vue3、Vite、Element Plus、Pinia、Vue Router和ECharts构建前端，相关技术生态成熟，能够满足表单录入、表格展示、路由控制、状态管理和图表可视化需求。后端采用Node.js和Koa2构建RESTful API，使用koa-router划分业务接口，使用koa-bodyparser解析请求体，使用koa-cors处理跨域访问。")
    para(doc, anchor, "数据层采用MongoDB和Mongoose实现文档存储与Schema约束，符合个人记账系统数据结构灵活、统计查询频繁、字段可扩展的特点。系统中已定义User、Transaction、Category、Account、UserConfig、Notification等模型，并在记账记录、分类、通知和风险等级等查询场景上设置索引，因此从技术选型和项目代码实现情况看，本系统具有可行性。")
    add_expand(doc, anchor, "feasibility")
    para(doc, anchor, "1.2.2 经济可行性", "Heading 4", False)
    para(doc, anchor, "系统使用的Vue、Koa、MongoDB、ECharts等主要技术均为开源技术，开发阶段无需额外授权费用。部署时前端静态资源可由Nginx托管，后端Node.js服务和MongoDB可部署在普通云服务器中，资源消耗较低，适合毕业设计规模和个人项目后续维护。")
    para(doc, anchor, "1.2.3 操作可行性", "Heading 4", False)
    para(doc, anchor, "系统采用Web应用形式，用户通过浏览器即可访问，不需要安装客户端。用户端以总览、交易记录、统计分析、消费分析、个性设置、消息中心等页面组织功能；管理员端以管理总览、用户列表、风险监控、消费统计、分类分析、系统设置、操作日志等页面组织功能，界面层级清晰，操作路径较短，具备较好的操作可行性。")
    para(doc, anchor, "1.3 相关技术", "Heading 3", False)
    para(doc, anchor, "前端部分使用Vue3作为核心框架，页面组件采用单文件组件方式组织。Vue Router负责用户端和管理员端路由切换，Pinia保存登录用户、权限状态和页面共享状态，Element Plus提供表单、表格、弹窗、菜单等基础组件，ECharts用于收支趋势、分类占比、风险分布等图表展示。")
    para(doc, anchor, "后端部分使用Node.js运行环境和Koa2框架。Koa的中间件机制适合组织鉴权、跨域、请求解析、错误处理等通用逻辑；jsonwebtoken用于生成和校验JWT；bcryptjs用于密码加密；xlsx用于解析微信和支付宝账单文件；node-cron用于定时检查预算预警和月度成就。")
    para(doc, anchor, "数据库部分使用MongoDB保存业务数据，并通过Mongoose定义Schema、默认值、枚举约束、索引和模型方法。由于本项目数据以用户、记账记录、通知、配置等文档为主，MongoDB能够较自然地表达嵌套配置和扩展字段。")
    add_expand(doc, anchor, "related")

    para(doc, anchor, "2 系统设计", "Heading 2", False)
    para(doc, anchor, "2.1 系统架构设计", "Heading 3", False)
    para(doc, anchor, "本系统采用前后端分离的B/S架构。前端负责页面展示、表单交互、路由跳转和图表渲染；后端负责用户鉴权、业务校验、数据统计、账单导入、预算预警和后台管理接口；数据层通过MongoDB保存业务集合。系统架构如图2-1所示。")
    pic(doc, anchor, paths["arch"], "图2-1 系统架构图", 15.0)
    para(doc, anchor, "从系统运行流程看，用户在浏览器中访问前端页面，前端通过Axios向后端RESTful API发送请求。后端在接收到请求后先完成JWT校验，再根据路由分发到认证、记账、账户、导入、通知或管理模块。业务处理完成后，后端通过Mongoose访问finance_db数据库，并将JSON结果返回给前端。")
    add_expand(doc, anchor, "architecture")
    para(doc, anchor, "2.2 系统功能模块设计", "Heading 3", False)
    para(doc, anchor, "根据项目代码中的前端路由、后端接口和实际页面功能，本系统功能模块不按菜单名称简单罗列，而按业务能力划分。用户端以个人收支数据为中心，重点体现记账、账单导入、收支趋势可视化、分类占比可视化、消费结构可视化和预算预警；管理端以系统数据治理为中心，重点体现用户权限、分类配置、全局统计可视化、分类消费分析可视化、风险监控可视化和日志导出。系统功能模块结构如图2-2所示。")
    pic(doc, anchor, paths["modules"], "图2-2 系统功能模块图", 15.5)
    para(doc, anchor, "用户端功能中，账号认证资料维护对应注册、登录、验证码校验、资料修改和密码修改；收支记账流水管理对应交易记录的新增、查询、编辑、删除和分页筛选；微信支付宝账单导入对应外部账单解析和批量写入；收支趋势可视化、分类占比可视化和消费结构可视化对应ECharts图表展示，是论文题目中“可视化分析”的主要体现；预算预警消息通知对应预算阈值检测、预警通知生成和消息读取。管理端功能中，用户权限管理、收支分类配置和系统参数模板管理负责基础数据维护；全局收支统计可视化、分类消费分析可视化和风险监控可视化负责管理员侧的数据观察；数据导出日志审计负责数据备份和后台操作追踪。")
    add_expand(doc, anchor, "modules")
    para(doc, anchor, "2.3 数据库设计", "Heading 3", False)
    para(doc, anchor, "本系统使用MongoDB数据库，数据库名称为finance_db。数据库集合根据Mongoose模型进行划分，主要包括users、transactions、accounts、categories、userconfigs、notifications、operationlogs、systemsettings、messagetemplates、verificationcodes等集合。")
    para(doc, anchor, "2.3.1 数据库概念设计", "Heading 4", False)
    para(doc, anchor, "系统核心业务围绕用户记账展开。用户可以创建多条记账记录，记账记录按照分类名称归入对应的收入或支出分类；账户、预算、主题、通知开关和风险评分等扩展信息在后续表结构设计中说明。核心业务E-R图如图2-3所示。")
    pic(doc, anchor, paths["core_er"], "图2-3 核心业务E-R图", 15.4)
    para(doc, anchor, "以用户创建一条记账记录为例，用户实体与记账记录实体之间为一对多关系，即一个用户可以创建多条记账记录，而每条记账记录只属于一个用户。用户记账实例E-R图如图2-4所示。")
    pic(doc, anchor, paths["inst_er"], "图2-4 用户记账实例E-R图", 14.6)
    para(doc, anchor, "E-R图中使用中文属性名表达业务含义，数据库表设计中使用项目代码中的真实字段名，两者在数量和含义上逐项对应。MongoDB实际文档中可能出现Mongoose自动生成的__v字段，该字段属于Mongoose的versionKey，用于内部版本控制，本系统业务代码未使用该字段，因此不作为数据库概念设计和表结构设计字段列出。")
    add_expand(doc, anchor, "database")
    para(doc, anchor, "2.3.2 数据库表设计", "Heading 4", False)
    para(doc, anchor, "数据库表结构依据server/models目录下的Mongoose模型整理，字段属性以Schema约束、索引和默认值为依据。")
    for idx, (name, rows) in enumerate(TABLES, 1):
        add_table(doc, anchor, f"表2-{idx} {name}", rows)

    para(doc, anchor, "3 系统详细设计与实现", "Heading 2", False)
    para(doc, anchor, "本章按照系统功能模块展开详细设计与实现说明。结合前文功能模块图和项目源码，系统详细设计不再按前端层、后端层、数据库层分别叙述，而是以用户实际使用路径为主线，将实现细节归入用户信息、记账管理、统计分析、预算通知和后台管理五个模块。每个模块下再说明具体功能的输入、处理过程、数据读写和页面展示方式，使论文结构与系统功能边界保持一致。")

    para(doc, anchor, "3.1 用户信息模块详细设计与实现", "Heading 3", False)
    para(doc, anchor, "用户信息模块是系统的入口模块，主要完成验证码发送、注册登录、用户资料修改和权限控制等功能。项目中该模块主要由auth路由、user路由、User模型、VerificationCode模型以及前端登录注册页面共同实现。该模块的设计重点在于确认用户身份、保护密码安全，并为后续记账数据隔离提供用户标识。")
    para(doc, anchor, "3.1.1 发送验证码功能", "Heading 4", False)
    para(doc, anchor, "发送验证码功能服务于用户注册和找回密码流程。用户在页面中输入邮箱后，前端将邮箱地址提交到后端验证码接口，后端生成验证码并保存到verificationcodes集合中。该集合包含email、code、type、expiresAt、createdAt等字段，其中expiresAt用于控制验证码有效时间。项目在VerificationCode模型中为expiresAt设置TTL索引，验证码过期后可由MongoDB自动清理，避免无效验证码长期留存在数据库中。")
    para(doc, anchor, "验证码校验时，后端会同时匹配邮箱、验证码和业务类型，确保注册验证码不能被错误地用于其他流程。验证码通过后，系统才继续执行用户注册或密码重置操作。这样的设计比只在前端判断验证码更可靠，因为真正的有效性判断发生在服务端，用户无法通过修改页面状态绕过校验。")
    para(doc, anchor, "3.1.2 登录注册功能", "Heading 4", False)
    para(doc, anchor, "注册功能需要用户提交用户名、邮箱、密码和验证码。后端首先检查必要字段是否完整，再查询验证码是否存在且未过期。验证码通过后，系统检查用户名是否已经被占用，并使用bcryptjs对密码进行哈希处理，最后创建User文档。User模型保存username、password、email、role、avatar、bio、createdAt、updatedAt等字段，其中password字段不直接保存明文密码。")
    para(doc, anchor, "登录功能以用户名和密码作为主要输入。后端根据username查询用户，查询到用户后使用bcrypt.compare比较明文密码与数据库中保存的哈希值。验证成功后，系统生成JWT并返回给前端，Token中包含用户id、用户名、邮箱和角色等必要信息。前端将Token保存到本地，并在后续接口请求中放入Authorization请求头。")
    add_expand(doc, anchor, "auth")
    para(doc, anchor, "3.1.3 用户资料修改功能", "Heading 4", False)
    para(doc, anchor, "用户资料修改功能主要用于维护昵称、邮箱、头像、个人简介和密码等信息。前端个人资料页面读取当前登录用户信息并填充表单，用户修改后提交到user或auth相关接口。后端通过Token解析出的用户id定位当前用户，只允许用户修改自己的资料，避免通过请求参数修改其他用户信息。")
    para(doc, anchor, "对于密码修改这类敏感操作，系统需要先验证旧密码或验证码，再写入新的哈希密码。头像、简介等展示信息则更偏向用户体验，更新成功后前端会同步刷新用户状态，使侧边栏、个人中心等位置展示最新资料。该功能虽然不直接影响记账流程，但它保证了用户账号的完整性和可维护性。")
    para(doc, anchor, "3.1.4 权限控制功能", "Heading 4", False)
    para(doc, anchor, "系统用户角色分为普通用户和管理员。普通用户进入用户端页面，管理员进入后台管理端页面。前端路由文件通过路由守卫判断Token和用户角色，普通用户访问/admin路径时会被重定向到用户端，管理员访问普通用户路径时会被引导到后台首页。")
    para(doc, anchor, "前端路由控制主要提升页面使用体验，真正的权限边界仍然放在后端接口。管理员接口在处理请求前会校验Token并检查role是否为admin；交易记录、账户、通知等用户数据接口则通过Token中的用户id过滤数据。这样既能防止普通用户访问管理员功能，也能防止用户读取或修改不属于自己的记账数据。")

    para(doc, anchor, "3.2 记账管理模块详细设计与实现", "Heading 3", False)
    para(doc, anchor, "记账管理模块是系统最核心的业务模块，围绕用户收支记录的新增、查询、修改、删除和账单导入展开。项目中该模块主要由transactions路由、import路由、Transaction模型、Account模型和用户端交易记录页面实现。交易记录是统计分析、预算预警和风险评估的基础，因此该模块需要保证数据录入准确、查询条件完整、用户数据严格隔离。")
    para(doc, anchor, "3.2.1 新增记账功能", "Heading 4", False)
    para(doc, anchor, "新增记账功能支持用户录入收入或支出记录。前端表单要求用户填写收支类型、金额、分类、日期和备注等信息，后端接收到请求后检查金额、类型和分类是否有效。Transaction模型中的主要字段包括userId、type、amount、category、note、date、createdAt和updatedAt。保存记录时，后端将Token解析出的用户id写入userId，确保每条交易记录都有明确归属。")
    para(doc, anchor, "项目中还支持批量新增交易记录，前端可以通过items数组一次提交多条记录。后端逐条检查金额、分类和类型，过滤掉缺少必要字段或分类类型不匹配的数据，再使用insertMany批量写入。相比前端循环调用新增接口，批量提交可以减少网络请求次数，也更适合账单导入后的集中保存场景。")
    para(doc, anchor, "3.2.2 交易记录查询功能", "Heading 4", False)
    para(doc, anchor, "交易记录查询功能服务于用户查看历史流水。前端交易记录页面提供收支类型、分类、日期范围和关键词等筛选条件，后端根据这些条件拼接MongoDB查询对象，并始终带上当前用户的userId。查询结果按照交易日期倒序排列，同时返回total、page、pageSize和totalPages等分页信息。")
    para(doc, anchor, "分页设计可以避免一次性加载全部历史交易记录。当用户持续使用系统后，transactions集合中的记录会不断增加，如果前端一次性请求全部数据，会造成页面渲染缓慢和网络传输浪费。通过分页和筛选，用户既可以快速浏览近期流水，也可以按条件查找某一类消费记录。")
    add_expand(doc, anchor, "transactions")
    if (SHOT_DIR / "图5-2.png").exists():
        pic(doc, anchor, SHOT_DIR / "图5-2.png", "图3-1 用户端交易记录页面", 14.8)
    para(doc, anchor, "3.2.3 交易编辑删除功能", "Heading 4", False)
    para(doc, anchor, "交易编辑功能用于修正用户录入错误的记录。用户在交易列表中选择某条记录后，前端打开编辑表单并回填原有数据；提交修改时，后端同时匹配记录_id和userId，只有记录属于当前登录用户时才允许更新。可修改内容包括收支类型、金额、分类、备注和交易日期等字段。")
    para(doc, anchor, "交易删除功能同样需要同时匹配_id和userId。这样设计可以避免用户通过猜测记录id删除其他用户的流水。删除成功后，前端刷新列表和统计数据，保证页面展示与数据库状态一致。交易编辑和删除虽然是常规功能，但它们直接关系到统计结果的准确性，因此必须保证权限校验和页面反馈都清晰可靠。")
    para(doc, anchor, "3.2.4 账单导入功能", "Heading 4", False)
    para(doc, anchor, "账单导入功能用于降低用户手工记账成本。系统支持微信和支付宝账单导入，后端使用xlsx库解析上传文件，并根据不同平台的表头结构识别交易时间、金额、收支类型、交易说明等信息。解析完成后，外部账单数据会被转换为系统统一的Transaction文档格式。")
    para(doc, anchor, "微信账单和支付宝账单的字段命名、表头位置和金额格式并不完全一致，因此导入逻辑不能简单按固定列号读取。项目代码会查找关键表头并进行字段映射，金额字段会去除人民币符号和其他非数字字符后再转换为数值。导入接口还支持accountId参数，当用户选择账户后，系统可以根据导入的收入和支出差额更新账户余额。")

    para(doc, anchor, "3.3 统计分析模块详细设计与实现", "Heading 3", False)
    para(doc, anchor, "统计分析模块负责将交易记录转化为可读的财务结果，主要包括首页总览、收支统计、分类占比分析和消费分析等功能。项目中该模块依赖transactions路由中的statistics、month-stats、daily-stats等接口，以及前端Dashboard、Statistics和ConsumptionAnalysis页面。该模块的设计目标不是单纯展示图表，而是帮助用户理解自己的收入、支出、结余和消费结构。")
    para(doc, anchor, "3.3.1 首页总览功能", "Heading 4", False)
    para(doc, anchor, "首页总览是用户登录后的主要入口，展示本月收入、本月支出、结余、预算使用情况和近期交易等信息。前端页面在加载时请求统计接口和交易列表接口，将聚合数据展示为概览卡片和趋势图。用户不需要进入多个页面即可快速判断当前财务状态。")
    para(doc, anchor, "首页总览功能强调信息密度和可读性。收入、支出和结余属于用户最关注的指标，因此放在页面显著位置；近期交易用于帮助用户核对最近流水；趋势图用于观察收支变化方向。该页面把记账数据、统计数据和预算信息组合在一起，是系统用户端的综合展示页面。")
    if (SHOT_DIR / "图5-1.png").exists():
        pic(doc, anchor, SHOT_DIR / "图5-1.png", "图3-2 用户端首页总览", 14.8)
    para(doc, anchor, "3.3.2 收支统计功能", "Heading 4", False)
    para(doc, anchor, "收支统计功能通过MongoDB聚合管道完成。后端按照当前用户id和日期范围筛选交易记录，再按type字段统计income和expense金额，并计算结余。对于每日趋势，系统按照交易日期分组，分别汇总每日收入和支出。聚合结果返回前端后，由ECharts渲染为折线图、柱状图或统计卡片。")
    para(doc, anchor, "后端集中计算统计结果可以保证统计口径一致。若把所有记录返回前端再计算，不仅增加网络传输量，也容易造成不同页面统计逻辑不一致。本系统将统计规则放在后端接口中，用户端首页、统计分析页面和部分管理员统计页面都可以复用类似的计算思路。")
    if (SHOT_DIR / "图5-3.png").exists():
        pic(doc, anchor, SHOT_DIR / "图5-3.png", "图3-3 用户端统计分析页面", 14.8)
    para(doc, anchor, "3.3.3 分类占比分析功能", "Heading 4", False)
    para(doc, anchor, "分类占比分析用于查看不同消费分类在总支出中的比例。系统根据category字段对支出记录进行分组，计算每个分类的金额和占比。分类字段虽然在Transaction中以字符串保存，但分类的可选项由categories集合维护，管理员可以通过后台分类管理保证收入分类和支出分类的规范性。")
    para(doc, anchor, "分类占比图能够帮助用户发现主要消费方向。例如餐饮、交通、购物等分类支出占比较高时，用户可以结合预算设置进行调整。该功能的准确性依赖于记账时分类选择的规范性，因此项目在新增记录时会校验分类与收支类型是否匹配，避免收入分类和支出分类混用。")
    para(doc, anchor, "3.3.4 消费分析功能", "Heading 4", False)
    para(doc, anchor, "消费分析功能在基础统计之上进一步解释用户消费习惯。前端ConsumptionAnalysis页面将支出趋势、分类结构和阶段性消费变化集中展示，使用户能够从时间和类别两个维度复盘消费行为。该功能面向的是长期使用场景，用户数据积累越多，分析结果越有参考价值。")
    para(doc, anchor, "从实现上看，消费分析仍以transactions集合为数据基础。系统按照用户、时间范围和支出类型筛选记录，再根据日期和分类进行聚合。前端负责把聚合结果转换为适合阅读的图表和文字提示。通过这种方式，系统把原始流水转化为可视化信息，符合论文摘要中提出的“记账与可视化分析”目标。")
    add_expand(doc, anchor, "frontend")
    if (SHOT_DIR / "图5-4.png").exists():
        pic(doc, anchor, SHOT_DIR / "图5-4.png", "图3-4 用户端消费分析页面", 14.8)

    para(doc, anchor, "3.4 预算通知模块详细设计与实现", "Heading 3", False)
    para(doc, anchor, "预算通知模块将用户配置、交易统计和消息提醒连接起来，主要包括预算设置、预算预警和消息通知三个功能。项目中该模块主要由UserConfig模型、Notification模型、notifications路由以及用户端个性设置和消息中心页面实现。该模块体现了系统从被动记录向主动提醒的扩展。")
    para(doc, anchor, "3.4.1 预算设置功能", "Heading 4", False)
    para(doc, anchor, "预算设置功能保存在userconfigs集合中。UserConfig模型包含budget.monthly、budget.yearly、budget.alertThreshold、notification.budgetAlert、notification.riskAlert、notification.systemNotice等配置字段。用户在个性设置页面修改预算金额和提醒阈值后，前端提交到后端配置接口，后端按照当前用户id更新对应配置。")
    para(doc, anchor, "预算字段与用户表分开保存，是因为预算、主题、通知开关和风险评分都属于用户扩展配置。如果全部放入users集合，会使用户基础信息和业务配置混在一起，不利于维护。单独的userconfigs集合可以让用户基础身份信息保持简洁，也便于后续扩展更多个性化设置。")
    para(doc, anchor, "3.4.2 预算预警功能", "Heading 4", False)
    para(doc, anchor, "预算预警功能在用户支出达到预设阈值时触发。系统读取用户月度预算和预警阈值，统计当月expense类型交易总额，并计算预算使用率。当使用率超过alertThreshold且用户开启预算提醒时，后端创建一条budget_alert类型通知。")
    para(doc, anchor, "项目中预算预警既可以在新增支出后触发，也可以通过定时任务周期性检查。为了避免重复提醒，系统会判断当天是否已经发送过同类预算预警，若已经发送则不再重复创建通知。该设计能够在提醒用户关注预算的同时，避免消息中心被重复预警打扰。")
    para(doc, anchor, "3.4.3 消息通知功能", "Heading 4", False)
    para(doc, anchor, "消息通知功能由notifications集合保存，主要字段包括userId、title、content、type、priority、isRead、relatedData、createdAt和updatedAt。通知类型包括预算预警、系统通知、风险提醒和交易相关通知。用户端消息中心可以查看通知列表、获取未读数量、标记已读、全部已读和删除通知。")
    para(doc, anchor, "通知功能使系统分析结果能够主动反馈给用户。预算预警、风险提醒和系统通知如果只停留在后台数据中，用户很难及时感知；通过消息中心集中展示，用户可以在使用系统时看到提醒并采取后续操作。该模块与预算配置、风险评估和管理员提醒功能都存在关联，是系统闭环的重要部分。")
    add_expand(doc, anchor, "risk")
    if (SHOT_DIR / "图5-6.png").exists():
        pic(doc, anchor, SHOT_DIR / "图5-6.png", "图3-5 用户端消息中心页面", 14.8)

    para(doc, anchor, "3.5 后台管理模块详细设计与实现", "Heading 3", False)
    para(doc, anchor, "后台管理模块面向管理员用户，主要完成系统数据概览、用户管理、分类管理、系统设置、风险监控、数据导出和操作日志查看等功能。项目中该模块主要由admin路由、管理员端页面、OperationLog模型、SystemSettings模型和MessageTemplate模型实现。后台管理模块与普通用户端共用finance_db数据库，但接口入口和权限校验不同。")
    para(doc, anchor, "3.5.1 管理端数据概览功能", "Heading 4", False)
    para(doc, anchor, "管理端数据概览页面通过dashboard-stats等接口汇总系统运行数据，包括用户总数、活跃用户、交易数量、交易金额和风险用户数量等指标。后端通过聚合查询users、transactions和userconfigs等集合生成概览结果，前端将其展示为统计卡片和趋势图。")
    para(doc, anchor, "该功能的作用是让管理员快速了解系统整体运行状态。对于后台管理端而言，首页不应只展示导航入口，而应优先呈现与管理决策相关的数据，例如用户增长、交易规模和风险用户数量。管理员可以根据概览数据进一步进入用户管理、风险监控或消费统计页面。")
    add_expand(doc, anchor, "admin")
    if (SHOT_DIR / "图5-7.png").exists():
        pic(doc, anchor, SHOT_DIR / "图5-7.png", "图3-6 管理端数据概览页面", 14.8)
    para(doc, anchor, "3.5.2 用户管理功能", "Heading 4", False)
    para(doc, anchor, "用户管理功能以用户列表为核心，支持分页查询、创建用户、编辑用户、删除用户和批量操作。管理员新增用户时，后端会检查用户名唯一性并创建User文档；编辑用户时可以维护用户名、邮箱、角色、预算和主题等信息；删除用户时需要谨慎处理，因为该操作会影响用户相关业务数据。")
    para(doc, anchor, "用户管理功能需要严格区分普通用户和管理员。普通用户只能维护自己的资料，而管理员可以从后台维护用户信息。项目通过管理员接口权限校验实现这一点，只有role为admin的用户才能访问admin路由下的用户管理接口。")
    para(doc, anchor, "3.5.3 分类管理与系统设置功能", "Heading 4", False)
    para(doc, anchor, "分类管理功能维护categories集合中的收入和支出分类，字段包括name、type、iconId、color、isDefault、sortOrder、createdAt和updatedAt。分类数据会影响用户新增记账时的可选项，也会影响分类统计页面的分析维度，因此分类名称、类型和排序都需要保持规范。")
    para(doc, anchor, "系统设置功能维护systemsettings集合中的系统名称、默认币种、默认预算和功能开关等配置。消息模板功能维护messagetemplates集合中的提醒文案，管理员发送风险提醒时可以复用模板内容。把这些配置保存在数据库中，可以减少修改代码的频率，也便于管理员通过页面维护系统运行参数。")
    para(doc, anchor, "3.5.4 风险监控功能", "Heading 4", False)
    para(doc, anchor, "风险监控功能基于UserConfig中的financialHealth字段和风险计算工具实现。风险算法从预算使用率、收支平衡和交易频率三个维度计算得分，并按照0.4、0.3、0.3的权重得到综合评分。综合评分低于40为高风险，40至70之间为中风险，70及以上为低风险。")
    para(doc, anchor, "管理员端风险监控页面展示不同风险等级用户，并提供筛选和提醒入口。管理员发现高风险或中风险用户后，可以发送风险提醒通知，通知最终写入notifications集合并在用户消息中心展示。该功能把算法评估结果转化为后台管理动作，使系统具备一定的主动风险干预能力。")
    if (SHOT_DIR / "图5-9.png").exists():
        pic(doc, anchor, SHOT_DIR / "图5-9.png", "图3-7 管理端风险监控页面", 14.8)
    para(doc, anchor, "3.5.5 数据导出与操作日志功能", "Heading 4", False)
    para(doc, anchor, "数据导出功能用于管理员备份和分析系统数据。admin路由中的export接口可以根据导出类型生成用户、交易或统计相关数据，并支持JSON、CSV、Excel等格式。对于记账系统而言，导出能力不仅服务管理员，也体现了数据可迁移性，避免业务数据只能留存在单一系统中。")
    para(doc, anchor, "操作日志功能记录管理员关键操作，operationlogs集合包含userId、action、details、createdAt和updatedAt等字段。日志可以帮助管理员追踪后台配置修改、用户操作和数据导出行为。虽然毕业设计版本的日志字段较为基础，但它已经体现了后台系统应具备的可追溯性，为后续审计和异常排查提供了数据基础。")

    para(doc, anchor, "4 系统测试", "Heading 2", False)
    para(doc, anchor, "4.1 测试目的", "Heading 3", False)
    para(doc, anchor, "系统测试的目的是验证个人财务记账与可视化分析系统是否满足需求分析中提出的功能需求和基本非功能需求。测试重点包括用户认证、记账管理、账单导入、统计分析、预算预警、消息通知、管理员管理和风险评估等模块。")
    add_expand(doc, anchor, "testing")
    para(doc, anchor, "4.2 功能测试", "Heading 3", False)
    test_table(doc, anchor, "表4-1 系统功能测试用例表", [
        ("TC01", "用户认证", "输入正确用户名和密码登录系统", "登录成功并进入对应用户端或管理端", "通过"),
        ("TC02", "用户认证", "输入错误密码登录系统", "系统提示用户名或密码错误", "通过"),
        ("TC03", "记账管理", "新增一条支出记录并填写金额、分类和日期", "记录保存成功并出现在交易列表中", "通过"),
        ("TC04", "记账管理", "按分类、类型、日期范围筛选交易记录", "列表仅显示满足条件的数据", "通过"),
        ("TC05", "账单导入", "上传微信或支付宝账单文件并执行导入", "系统解析有效记录并返回导入数量", "通过"),
        ("TC06", "统计分析", "查看收入、支出、结余和分类占比", "图表和汇总金额按交易数据正确计算", "通过"),
        ("TC07", "预算预警", "支出达到预算预警阈值", "系统生成预算预警通知", "通过"),
        ("TC08", "消息通知", "将未读消息标记为已读", "未读数量减少，消息状态更新", "通过"),
        ("TC09", "后台管理", "管理员新增或修改系统分类", "分类保存成功并在分类列表展示", "通过"),
        ("TC10", "风险监控", "管理员查看风险用户列表", "系统按风险等级展示用户及评分", "通过"),
    ])
    para(doc, anchor, "4.3 性能与兼容性测试", "Heading 3", False)
    test_table(doc, anchor, "表4-2 系统性能与兼容性测试表", [
        ("PT01", "页面加载", "首次进入用户端首页", "页面在可接受时间内完成主要内容渲染", "通过"),
        ("PT02", "分页查询", "交易记录分页查询20条数据", "接口返回分页数据和总数", "通过"),
        ("PT03", "聚合统计", "查询指定时间范围内统计数据", "后端通过聚合管道返回统计结果", "通过"),
        ("PT04", "权限控制", "普通用户直接访问/admin路径", "系统重定向到普通用户页面", "通过"),
        ("PT05", "浏览器兼容", "在现代浏览器中访问主要页面", "页面布局正常，图表显示完整", "通过"),
    ])
    para(doc, anchor, "4.4 测试分析", "Heading 3", False)
    para(doc, anchor, "从测试结果看，系统核心功能能够按照设计要求运行。普通用户可以完成注册登录、记账、查询、统计、预算配置和消息查看，管理员可以完成用户管理、分类管理、风险监控和数据导出等操作。后端接口能够根据Token区分用户身份，数据库查询能够按用户id隔离数据，统计接口能够返回收入、支出、结余和分类聚合结果。")
    para(doc, anchor, "测试过程中也反映出系统仍有进一步优化空间，例如导入文件格式的兼容范围可以继续扩大，统计页面在大量交易记录下可增加缓存策略，管理员批量操作可以补充更细粒度的操作确认。总体而言，系统已达到毕业设计阶段的功能目标和基本质量要求。")
    add_expand(doc, anchor, "test_analysis")

    para(doc, anchor, "5 研究或成果与社会、环境的关系", "Heading 2", False)
    para(doc, anchor, "本系统以个人财务记录和可视化分析为研究对象，能够帮助用户将分散在移动支付平台中的消费数据转化为可查询、可统计、可预警的个人财务信息。从社会价值看，系统通过预算预警、消费分析和风险评估，引导用户关注自身收支结构，减少冲动消费和无计划支出，有助于提升个人和家庭的财务管理意识。")
    para(doc, anchor, "从教育意义看，系统将前端工程化、后端RESTful接口、MongoDB文档模型、JWT认证、图表可视化和定时任务等技术结合在一个完整项目中，体现了软件工程中需求分析、系统设计、编码实现和测试验证的基本流程，对Web应用开发实践具有一定参考价值。")
    para(doc, anchor, "从环境关系看，系统采用Web应用方式提供服务，用户无需纸质账本记录收支，也无需频繁安装和更新本地客户端，有助于减少纸张消耗和软件分发成本。系统通过前后端分离和模块化设计支持后续按需扩展，能够在用户规模增长时逐步优化部署资源，避免一次性投入过高的服务器资源。")
    para(doc, anchor, "当然，系统在产生积极作用的同时也需要重视数据安全和隐私保护。个人财务数据具有较强敏感性，因此系统通过密码加密、Token鉴权、用户数据隔离和管理员权限控制降低数据泄露风险。后续若投入真实使用，还需要进一步完善传输加密、备份恢复、异常审计和隐私授权说明。")
    add_expand(doc, anchor, "social")


def main():
    paths = {
        "arch": save_architecture(),
        "modules": save_modules(),
        "core_er": save_core_er(),
        "inst_er": save_instance_er(),
    }
    doc = Document(SRC)
    anchor = delete_allowed_block(doc)
    build_content(doc, anchor, paths)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
