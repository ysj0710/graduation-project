from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/enjoy0710/Desktop/学校相关/毕业设计/graduation-project")
OUT = ROOT / "论文图表补充"
OUT.mkdir(exist_ok=True)

FONT_PATH = "/System/Library/Fonts/STHeiti Medium.ttc"


def font(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_PATH, size=size, index=0)


F16 = font(16)
F14 = font(14)
F15 = font(15)
F18 = font(18)
F20 = font(20)
F22 = font(22)
F24 = font(24)
F28 = font(28)


def text_center(draw, box, text, fnt=F18, fill=(20, 20, 20)):
    x1, y1, x2, y2 = box
    lines = str(text).split("\n")
    heights = []
    widths = []
    for line in lines:
        b = draw.textbbox((0, 0), line, font=fnt)
        widths.append(b[2] - b[0])
        heights.append(b[3] - b[1])
    total_h = sum(heights) + (len(lines) - 1) * 4
    y = y1 + (y2 - y1 - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + 4


def rect(draw, box, text="", fnt=F18, width=2, fill=(255, 255, 255), outline=(25, 25, 25)):
    draw.rectangle(box, fill=fill, outline=outline, width=width)
    if text:
        text_center(draw, box, text, fnt)


def line(draw, p1, p2, width=2):
    draw.line([p1, p2], fill=(25, 25, 25), width=width)


def ellipse(draw, box, text, fnt=F14 if False else F16):
    draw.ellipse(box, fill=(255, 255, 255), outline=(25, 25, 25), width=2)
    text_center(draw, box, text, fnt)


def diamond(draw, center, w, h, text, fnt=F16):
    x, y = center
    pts = [(x, y - h // 2), (x + w // 2, y), (x, y + h // 2), (x - w // 2, y)]
    draw.polygon(pts, fill=(255, 255, 255), outline=(25, 25, 25))
    draw.line(pts + [pts[0]], fill=(25, 25, 25), width=2)
    text_center(draw, (x - w // 2, y - h // 2, x + w // 2, y + h // 2), text, fnt)


def save_architecture():
    img = Image.new("RGB", (1200, 840), "white")
    d = ImageDraw.Draw(img)
    margin = 28
    rows = [
        (40, 150, "表示层 UI", ["Vue 3", "Element Plus", "ECharts", "Pinia", "Axios"]),
        (170, 430, "业务逻辑层 BLL", ["身份认证", "记账管理", "账户管理", "分类管理", "预算预警", "消息通知", "风险统计", "导入解析"]),
        (450, 580, "服务支撑层", ["Koa 2", "JWT", "Mongoose", "node-cron", "xlsx", "nodemailer"]),
        (600, 700, "数据交互层 DAL", ["MongoDB", "finance_db", "Mongoose Schema", "索引优化"]),
        (720, 805, "部署环境", ["Nginx", "云服务器", "HTTPS", "Node.js"]),
    ]
    for y1, y2, title, items in rows:
        rect(d, (margin, y1, 1172, y2), "", F18, width=2)
        rect(d, (margin, y1, 220, y2), title, F22, width=2, fill=(246, 248, 250))
        cols = 4 if len(items) <= 4 else 5
        gap = 22
        start_x = 250
        area_w = 900
        box_w = (area_w - gap * (cols - 1)) // cols
        box_h = 52
        for i, item in enumerate(items):
            cx = i % cols
            cy = i // cols
            x = start_x + cx * (box_w + gap)
            y = y1 + 25 + cy * 68
            rect(d, (x, y, x + box_w, y + box_h), item, F18, width=2)
    path = OUT / "图4-1 系统架构图.png"
    img.save(path)
    return path


def save_function_modules():
    img = Image.new("RGB", (1400, 650), "white")
    d = ImageDraw.Draw(img)
    rect(d, (420, 30, 980, 90), "个人财务记账系统", F24, width=2)
    modules = [
        ("用户信息模块", ["用户注册", "用户登录", "个人资料", "主题配置"]),
        ("记账管理模块", ["新增记账", "批量导入", "流水查询", "分类筛选", "收支统计"]),
        ("账户分类模块", ["账户维护", "账户余额", "分类维护", "图标颜色"]),
        ("风险预警模块", ["预算设置", "消费分析", "风险评分", "预警提醒"]),
        ("后台管理模块", ["用户管理", "分类管理", "数据看板", "操作日志", "系统设置"]),
        ("消息通知模块", ["通知发送", "通知读取", "模板配置", "定时任务"]),
    ]
    x0 = 40
    w = 210
    gap = 18
    for mi, (name, funcs) in enumerate(modules):
        x = x0 + mi * (w + gap)
        rect(d, (x, 140, x + w, 195), name, F18, width=2)
        line(d, (700, 90), (x + w // 2, 140), width=2)
        for j, f in enumerate(funcs):
            bw = 42
            bx = x + 10 + j * ((w - 20) / len(funcs))
            rect(d, (int(bx), 225, int(bx + bw), 570), f.replace("", "\n")[1:-1], F16, width=2)
            line(d, (x + w // 2, 195), (int(bx + bw // 2), 225), width=2)
    path = OUT / "图4-2 系统功能模块图.png"
    img.save(path)
    return path


ENTITY_FIELDS = {
    "用户": ["_id", "username", "password", "email", "role", "nickname", "avatar", "createdAt", "updatedAt", "__v"],
    "账户": ["_id", "userId", "name", "type", "icon", "color", "balance", "remark", "createdAt", "updatedAt", "__v"],
    "记账记录": ["_id", "userId", "type", "amount", "category", "note", "date", "isDeleted", "createdAt", "updatedAt", "__v"],
    "分类": ["_id", "name", "iconId", "color", "type", "isDefault", "sortOrder", "createdAt", "updatedAt", "__v"],
    "用户配置": ["_id", "userId", "budget.monthly", "budget.yearly", "budget.alertThreshold", "theme.background", "theme.primaryColor", "theme.glassBlur", "theme.pattern", "theme.presetId", "theme.customBgUrl", "notification.budgetAlert", "notification.riskAlert", "notification.systemAlert", "financialHealth.score", "financialHealth.riskLevel", "financialHealth.lastCalculatedAt", "createdAt", "updatedAt", "__v"],
    "通知": ["_id", "userId", "senderId", "type", "title", "content", "scope", "targetUserIds", "isRead", "data", "createdAt", "updatedAt", "__v"],
    "操作日志": ["_id", "userId", "action", "details", "createdAt", "updatedAt", "__v"],
    "系统设置": ["_id", "systemName", "currency", "defaultBudget", "createdAt", "updatedAt", "__v"],
    "消息模板": ["_id", "low", "medium", "high", "createdAt", "updatedAt", "__v"],
    "验证码": ["_id", "email", "code", "type", "expiresAt", "createdAt", "__v"],
}


def draw_entity_with_attrs(d, center, name, attrs, radius=300):
    x, y = center
    rect(d, (x - 58, y - 30, x + 58, y + 30), name, F18, width=2)
    n = len(attrs)
    for i, attr in enumerate(attrs):
        ang = -math.pi / 2 + 2 * math.pi * i / n
        ex = x + math.cos(ang) * radius
        ey = y + math.sin(ang) * radius * 0.72
        box = (int(ex - 62), int(ey - 24), int(ex + 62), int(ey + 24))
        line(d, (x, y), (int(ex), int(ey)), width=1)
        ellipse(d, box, attr, F14 if False else F16)


def draw_chen_entity(d, center, name, attrs, positions, rect_size=(120, 58), fnt=F14):
    x, y = center
    rw, rh = rect_size
    rect_box = (x - rw // 2, y - rh // 2, x + rw // 2, y + rh // 2)
    rect(d, rect_box, name, F18, width=2)
    for attr, pos in zip(attrs, positions):
        ax, ay = pos
        ew = max(96, min(180, len(attr) * 12 + 36))
        eh = 46 if len(attr) < 14 else 54
        attr_box = (int(ax - ew // 2), int(ay - eh // 2), int(ax + ew // 2), int(ay + eh // 2))
        line(d, (x, y), (ax, ay), width=1)
        ellipse(d, attr_box, attr, fnt)
    return rect_box


def point_on_rect(center, target, rect_size):
    cx, cy = center
    tx, ty = target
    rw, rh = rect_size[0] / 2, rect_size[1] / 2
    dx, dy = tx - cx, ty - cy
    if dx == 0 and dy == 0:
        return int(cx), int(cy)
    scale = min(rw / abs(dx) if dx else 10**9, rh / abs(dy) if dy else 10**9)
    return int(cx + dx * scale), int(cy + dy * scale)


def attr_line(d, center, pos, rect_size):
    line(d, point_on_rect(center, pos, rect_size), pos, width=1)


def draw_chen_entity_clean(d, center, name, attrs, positions, rect_size=(120, 58), fnt=F16):
    x, y = center
    rw, rh = rect_size
    rect_box = (x - rw // 2, y - rh // 2, x + rw // 2, y + rh // 2)
    rect(d, rect_box, name, F18, width=2)
    for attr, pos in zip(attrs, positions):
        ax, ay = pos
        ew = max(88, min(176, len(attr) * 12 + 34))
        eh = 46 if len(attr) <= 13 else 54
        attr_line(d, center, pos, rect_size)
        ellipse(d, (int(ax - ew // 2), int(ay - eh // 2), int(ax + ew // 2), int(ay + eh // 2)), attr, fnt)
    return rect_box


def relation_clean(d, c1, c2, label, a, b, rect1, rect2, diamond_size=(110, 70), offset=(0, 0), fnt=F16):
    mx, my = (c1[0] + c2[0]) // 2 + offset[0], (c1[1] + c2[1]) // 2 + offset[1]
    p1 = point_on_rect(c1, (mx, my), rect1)
    p2 = point_on_rect(c2, (mx, my), rect2)
    line(d, p1, (mx - diamond_size[0] // 2, my), width=2)
    line(d, (mx + diamond_size[0] // 2, my), p2, width=2)
    diamond(d, (mx, my), diamond_size[0], diamond_size[1], label, fnt)
    d.text((p1[0] + (mx - p1[0]) * 0.38, p1[1] + (my - p1[1]) * 0.38 - 22), a, font=F18, fill=(20, 20, 20))
    d.text((p2[0] + (mx - p2[0]) * 0.38, p2[1] + (my - p2[1]) * 0.38 - 22), b, font=F18, fill=(20, 20, 20))


def diamond_anchor(center, from_point, size):
    cx, cy = center
    fw, fh = size[0] / 2, size[1] / 2
    dx, dy = from_point[0] - cx, from_point[1] - cy
    if dx == 0 and dy == 0:
        return int(cx), int(cy)
    denom = abs(dx) / fw + abs(dy) / fh
    return int(cx + dx / denom), int(cy + dy / denom)


def relation_at(d, c1, c2, rel_c, label, a, b, rect1, rect2, diamond_size=(120, 76), fnt=F18):
    p1 = point_on_rect(c1, rel_c, rect1)
    p2 = point_on_rect(c2, rel_c, rect2)
    q1 = diamond_anchor(rel_c, p1, diamond_size)
    q2 = diamond_anchor(rel_c, p2, diamond_size)
    line(d, p1, q1, width=2)
    line(d, q2, p2, width=2)
    diamond(d, rel_c, diamond_size[0], diamond_size[1], label, fnt)
    d.text((p1[0] + (q1[0] - p1[0]) * 0.35, p1[1] + (q1[1] - p1[1]) * 0.35 - 24), a, font=F22, fill=(20, 20, 20))
    d.text((p2[0] + (q2[0] - p2[0]) * 0.35, p2[1] + (q2[1] - p2[1]) * 0.35 - 24), b, font=F22, fill=(20, 20, 20))


def oval_positions(cx, cy, n, rx=260, ry=170, start=-math.pi / 2):
    return [
        (
            int(cx + math.cos(start + 2 * math.pi * i / n) * rx),
            int(cy + math.sin(start + 2 * math.pi * i / n) * ry),
        )
        for i in range(n)
    ]


def side_positions(cx, cy, n, left=True, gap=58):
    x = cx - 245 if left else cx + 245
    y0 = cy - (n - 1) * gap // 2
    return [(x, y0 + i * gap) for i in range(n)]


def grid_positions(points):
    return [(int(x), int(y)) for x, y in points]


def save_er_overall():
    img = Image.new("RGB", (1800, 1250), "white")
    d = ImageDraw.Draw(img)
    user_c = (430, 790)
    tx_c = (900, 210)
    cat_c = (1370, 790)
    user_rect = (120, 60)
    tx_rect = (138, 60)
    cat_rect = (120, 60)
    user_pos = grid_positions([
        (430, 560), (245, 500), (120, 640), (110, 800), (160, 965),
        (300, 1110), (475, 1160), (650, 1110), (770, 955), (770, 665)
    ])
    tx_pos = grid_positions([
        (900, 45), (720, 55), (1080, 55), (610, 145), (1190, 145),
        (555, 250), (1245, 250), (505, 360), (1295, 360), (485, 460), (1315, 460)
    ])
    cat_pos = grid_positions([
        (1370, 560), (1555, 500), (1680, 640), (1690, 800), (1640, 965),
        (1500, 1110), (1325, 1160), (1150, 1110), (1030, 955), (1030, 665)
    ])
    draw_chen_entity_clean(d, user_c, "用户", ENTITY_FIELDS["用户"], user_pos, rect_size=user_rect, fnt=F16)
    draw_chen_entity_clean(d, tx_c, "记账记录", ENTITY_FIELDS["记账记录"], tx_pos, rect_size=tx_rect, fnt=F16)
    draw_chen_entity_clean(d, cat_c, "分类", ENTITY_FIELDS["分类"], cat_pos, rect_size=cat_rect, fnt=F16)
    relation_at(d, user_c, tx_c, (650, 520), "记录", "1", "n", user_rect, tx_rect, diamond_size=(124, 78))
    relation_at(d, tx_c, cat_c, (1150, 520), "归属", "n", "1", tx_rect, cat_rect, diamond_size=(124, 78))
    path = OUT / "图4-3 总体ER图.png"
    img.save(path)
    return path


def save_er_management():
    img = Image.new("RGB", (1800, 1200), "white")
    d = ImageDraw.Draw(img)
    config_c = (900, 610)
    config_rect = (138, 60)
    config_pos = grid_positions([
        (900, 70), (730, 95), (1070, 95), (565, 170), (1235, 170),
        (420, 285), (1380, 285), (330, 430), (1470, 430), (300, 610),
        (1500, 610), (330, 790), (1470, 790), (420, 935), (1380, 935),
        (565, 1050), (1235, 1050), (730, 1125), (900, 1140), (1070, 1125)
    ])
    draw_chen_entity_clean(d, config_c, "用户配置", ENTITY_FIELDS["用户配置"], grid_positions([
        *config_pos
    ]), rect_size=config_rect, fnt=F14)
    path = OUT / "图4-4 系统管理ER图.png"
    img.save(path)
    return path


def save_er_instance():
    img = Image.new("RGB", (1500, 980), "white")
    d = ImageDraw.Draw(img)
    user_c = (380, 500)
    tx_c = (1120, 500)
    user_pos = grid_positions([
        (380, 190), (190, 240), (105, 370), (100, 520), (130, 680),
        (260, 820), (420, 850), (570, 770), (620, 630), (610, 350)
    ])
    tx_pos = grid_positions([
        (1120, 185), (930, 230), (1310, 230), (1420, 370), (1420, 520),
        (1380, 680), (1260, 820), (1080, 850), (930, 770), (890, 640), (930, 360)
    ])
    user_rect = (120, 60)
    tx_rect = (138, 60)
    draw_chen_entity_clean(d, user_c, "用户", ENTITY_FIELDS["用户"], user_pos, rect_size=user_rect, fnt=F16)
    draw_chen_entity_clean(d, tx_c, "记账记录", ENTITY_FIELDS["记账记录"], tx_pos, rect_size=tx_rect, fnt=F16)
    relation_at(d, user_c, tx_c, (755, 500), "记录", "1", "n", user_rect, tx_rect, diamond_size=(135, 84), fnt=F18)
    path = OUT / "图4-5 用户记账实例ER图.png"
    img.save(path)
    return path


FIELDS = {
    "users": [
        ("_id", "ObjectId", "是", "否", "主键"),
        ("username", "String", "否", "否", "登录用户名，唯一，3-20位"),
        ("password", "String", "否", "否", "加密后的登录密码"),
        ("email", "String", "否", "是", "邮箱，唯一稀疏索引"),
        ("role", "String", "否", "否", "admin/user，默认 user"),
        ("nickname", "String", "否", "是", "用户昵称"),
        ("avatar", "String", "否", "是", "头像地址或 Base64 数据"),
        ("createdAt", "Date", "否", "否", "创建时间"),
        ("updatedAt", "Date", "否", "否", "更新时间"),
        ("__v", "Int32", "否", "否", "Mongoose 版本字段"),
    ],
    "accounts": [
        ("_id", "ObjectId", "是", "否", "主键"),
        ("userId", "ObjectId", "否", "否", "关联 users._id"),
        ("name", "String", "否", "否", "账户名称"),
        ("type", "String", "否", "否", "cash/bank/alipay/wechat/other"),
        ("icon", "String", "否", "是", "账户图标"),
        ("color", "String", "否", "是", "账户颜色"),
        ("balance", "Number", "否", "否", "账户余额"),
        ("remark", "String", "否", "是", "备注"),
        ("createdAt", "Date", "否", "否", "创建时间"),
        ("updatedAt", "Date", "否", "否", "更新时间"),
        ("__v", "Int32", "否", "否", "Mongoose 版本字段"),
    ],
    "transactions": [
        ("_id", "ObjectId", "是", "否", "主键"),
        ("userId", "ObjectId", "否", "否", "关联 users._id"),
        ("type", "String", "否", "否", "income/expense"),
        ("amount", "Number", "否", "否", "交易金额，最小值 0"),
        ("category", "String", "否", "否", "分类名称，对应 categories.name"),
        ("note", "String", "否", "是", "交易备注"),
        ("date", "Date", "否", "否", "交易发生时间"),
        ("isDeleted", "Boolean", "否", "是", "实际库中存在的软删除标识"),
        ("createdAt", "Date", "否", "否", "创建时间"),
        ("updatedAt", "Date", "否", "否", "更新时间"),
        ("__v", "Int32", "否", "否", "Mongoose 版本字段"),
    ],
    "categories": [
        ("_id", "ObjectId", "是", "否", "主键"),
        ("name", "String", "否", "否", "分类名称，唯一"),
        ("iconId", "String", "否", "是", "图标标识"),
        ("color", "String", "否", "否", "分类颜色"),
        ("type", "String", "否", "否", "income/expense"),
        ("isDefault", "Boolean", "否", "否", "是否默认分类"),
        ("sortOrder", "Number", "否", "否", "排序值"),
        ("createdAt", "Date", "否", "否", "创建时间"),
        ("updatedAt", "Date", "否", "否", "更新时间"),
        ("__v", "Int32", "否", "否", "Mongoose 版本字段"),
    ],
    "userconfigs": [
        ("_id", "ObjectId", "是", "否", "主键"),
        ("userId", "ObjectId", "否", "否", "关联 users._id，唯一"),
        ("budget.monthly", "Number", "否", "否", "月预算"),
        ("budget.yearly", "Number", "否", "否", "年预算"),
        ("budget.alertThreshold", "Number", "否", "否", "预算预警阈值"),
        ("theme.background", "String", "否", "否", "背景样式"),
        ("theme.primaryColor", "String", "否", "否", "主题主色"),
        ("theme.glassBlur", "Number", "否", "否", "玻璃模糊度"),
        ("theme.pattern", "String", "否", "否", "背景图案"),
        ("theme.presetId", "String", "否", "否", "主题预设 ID"),
        ("theme.customBgUrl", "String", "否", "是", "自定义背景地址"),
        ("notification.budgetAlert", "Boolean", "否", "否", "预算通知开关"),
        ("notification.riskAlert", "Boolean", "否", "否", "风险通知开关"),
        ("notification.systemAlert", "Boolean", "否", "否", "系统通知开关"),
        ("financialHealth.score", "Number", "否", "否", "财务健康评分"),
        ("financialHealth.riskLevel", "String", "否", "否", "low/medium/high"),
        ("financialHealth.lastCalculatedAt", "Date", "否", "否", "最近计算时间"),
        ("createdAt", "Date", "否", "否", "创建时间"),
        ("updatedAt", "Date", "否", "否", "更新时间"),
        ("__v", "Int32", "否", "否", "Mongoose 版本字段"),
    ],
    "notifications": [
        ("_id", "ObjectId", "是", "否", "主键"),
        ("userId", "ObjectId", "否", "否", "接收用户，关联 users._id"),
        ("senderId", "ObjectId", "否", "是", "发送者，关联 users._id"),
        ("type", "String", "否", "否", "budget_alert/system/risk_warning/transaction"),
        ("title", "String", "否", "否", "通知标题"),
        ("content", "String", "否", "否", "通知内容"),
        ("scope", "String", "否", "否", "all/selected"),
        ("targetUserIds", "Array<ObjectId>", "否", "是", "目标用户列表"),
        ("isRead", "Boolean", "否", "否", "是否已读"),
        ("data", "Mixed", "否", "是", "扩展数据"),
        ("createdAt", "Date", "否", "否", "创建时间"),
        ("updatedAt", "Date", "否", "否", "更新时间"),
        ("__v", "Int32", "否", "否", "Mongoose 版本字段"),
    ],
    "operationlogs": [
        ("_id", "ObjectId", "是", "否", "主键"),
        ("userId", "ObjectId", "否", "否", "操作用户，关联 users._id"),
        ("action", "String", "否", "否", "操作名称"),
        ("details", "String", "否", "是", "操作详情"),
        ("createdAt", "Date", "否", "否", "创建时间"),
        ("updatedAt", "Date", "否", "否", "更新时间"),
        ("__v", "Int32", "否", "否", "Mongoose 版本字段"),
    ],
    "systemsettings": [
        ("_id", "ObjectId", "是", "否", "主键"),
        ("systemName", "String", "否", "否", "系统名称"),
        ("currency", "String", "否", "否", "默认币种"),
        ("defaultBudget", "Number", "否", "否", "默认预算"),
        ("createdAt", "Date", "否", "否", "创建时间"),
        ("updatedAt", "Date", "否", "否", "更新时间"),
        ("__v", "Int32", "否", "否", "Mongoose 版本字段"),
    ],
    "messagetemplates": [
        ("_id", "ObjectId", "是", "否", "主键"),
        ("low", "String", "否", "否", "低风险提示模板"),
        ("medium", "String", "否", "否", "中风险提示模板"),
        ("high", "String", "否", "否", "高风险提示模板"),
        ("createdAt", "Date", "否", "否", "创建时间"),
        ("updatedAt", "Date", "否", "否", "更新时间"),
        ("__v", "Int32", "否", "否", "Mongoose 版本字段"),
    ],
    "verificationcodes": [
        ("_id", "ObjectId", "是", "否", "主键"),
        ("email", "String", "否", "否", "邮箱"),
        ("code", "String", "否", "否", "验证码"),
        ("type", "String", "否", "否", "register/resetPassword"),
        ("expiresAt", "Date", "否", "否", "过期时间，TTL 索引"),
        ("createdAt", "Date", "否", "否", "创建时间"),
        ("__v", "Int32", "否", "否", "Mongoose 版本字段"),
    ],
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def style_doc(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        st = styles[style_name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10.5)


def add_picture_with_caption(doc, img_path: Path, caption: str, width_cm: float = 16.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_table(doc, caption, rows):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["字段名", "数据类型", "主键", "允许空", "说明"]
    widths = [4.5, 3.2, 1.5, 1.8, 8.0]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "D9EAF7")
        set_cell_width(cell, widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            set_cell_width(cells[i], widths[i])
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(9)
    doc.add_paragraph()


def build_doc(paths):
    doc = Document()
    style_doc(doc)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("个人财务记账系统论文图表补充")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    p = doc.add_paragraph("以下图表依据毕业设计项目源码、Mongoose 模型及本机 finance_db 数据库结构整理。ER 图属性与数据库设计表字段保持一致；嵌套文档字段采用点路径表示。")
    p.paragraph_format.first_line_indent = Cm(0.74)

    add_picture_with_caption(doc, paths["arch"], "图4-1 系统架构图", 16.5)
    add_picture_with_caption(doc, paths["modules"], "图4-2 系统功能模块图", 16.5)
    doc.add_page_break()
    add_picture_with_caption(doc, paths["er"], "图4-3 核心业务ER图", 17.0)
    add_picture_with_caption(doc, paths["er_mgmt"], "图4-4 系统管理ER图", 17.0)
    add_picture_with_caption(doc, paths["er_inst"], "图4-5 用户记账实例ER图", 16.5)
    doc.add_page_break()
    screenshot = ROOT / "system_homepage_screenshot.png"
    if screenshot.exists():
        add_picture_with_caption(doc, screenshot, "图5-1 系统登录界面截图", 8.5)
    shots_dir = OUT / "功能截图"
    screenshots = [
        ("图5-2 用户端首页总览", "图5-1.png"),
        ("图5-3 用户端交易记录", "图5-2.png"),
        ("图5-4 用户端统计分析", "图5-3.png"),
        ("图5-5 用户端消费分析", "图5-4.png"),
        ("图5-6 用户端个性设置", "图5-5.png"),
        ("图5-7 用户端消息中心", "图5-6.png"),
        ("图5-8 管理端数据概览", "图5-7.png"),
        ("图5-9 管理端用户列表", "图5-8.png"),
        ("图5-10 管理端风险监控", "图5-9.png"),
        ("图5-11 管理端消费统计", "图5-10.png"),
        ("图5-12 管理端分类分析", "图5-11.png"),
    ]
    for caption, filename in screenshots:
        p = shots_dir / filename
        if p.exists():
            add_picture_with_caption(doc, p, caption, 16.5)
    doc.add_page_break()

    table_no = 1
    names = {
        "users": "用户表",
        "accounts": "账户表",
        "transactions": "记账记录表",
        "categories": "分类表",
        "userconfigs": "用户配置表",
        "notifications": "通知表",
        "operationlogs": "操作日志表",
        "systemsettings": "系统设置表",
        "messagetemplates": "消息模板表",
        "verificationcodes": "验证码表",
    }
    for key, rows in FIELDS.items():
        add_table(doc, f"表4-{table_no} {names[key]}（{key}）结构设计", rows)
        table_no += 1

    out_docx = ROOT / "毕业设计图表与数据库设计补充.docx"
    doc.save(out_docx)
    return out_docx


if __name__ == "__main__":
    paths = {
        "arch": save_architecture(),
        "modules": save_function_modules(),
        "er": save_er_overall(),
        "er_mgmt": save_er_management(),
        "er_inst": save_er_instance(),
    }
    docx = build_doc(paths)
    print(docx)
